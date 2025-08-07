# app.py

import dash
from dash import html
from dash.dependencies import Input, Output

from styles import external_stylesheets, nav_style
from components import nav_content
from components2 import nav_content2,nav_content3
from dash import dcc


import pandas as pd

from collections import Counter
import numpy as np
from dash import dcc, html
from dash import Dash, dcc, html, dash_table, Input, Output, State, callback, ctx



from data_processing import DataPreprocessing
#from EDA import eda
import base64
import io
import plotly.graph_objs as go
from input import page_5_layout
from filter import page_6_layout
from image import page_7_layout
from excel_to_word import page_10_layout
from docx import Document




import warnings
warnings.filterwarnings('ignore') 


@callback(
	Output("download-dataframe-csv3", "data"),
	Input('stored-data1', 'data'),
	Input('sped_button', 'n_clicks'),

	)

def save_data(data1, n_clicks):
	if(n_clicks != None):
		df = pd.DataFrame(data1)
		#df.to_csv("sped_data.csv", index=False)
		return dcc.send_data_frame(df.to_csv, "sped_data.csv", index=False)

@callback(
	Output("download-dataframe-pdf", "data"),
	Input('stored_results', 'data'),
	Input('sped_button8', 'n_clicks'),

	)

def save_data(data1, n_clicks):
	if(n_clicks != None):
		df = pd.DataFrame(data1)
		#df.to_csv("sped_data.csv", index=False)
		return dcc.send_data_frame(df.to_csv, "pdf_details.csv", index=False)


@callback(
	Output("download-dataframe-csv6", "data"),
	
	Input('stored-data-6f', 'data'),
	Input('sped_button6', 'n_clicks'),

	)

def save_data(data1, n_clicks):
	if(n_clicks != None):
		df = pd.DataFrame(data1)
		#df.to_csv("sped_data.csv", index=False)
		return dcc.send_data_frame(df.to_csv, "allzeros.csv", index=False)

@callback(
	Output("download-dataframe-csv7", "data"),
	
	Input('stored-data-7f', 'data'),
	Input('sped_button7', 'n_clicks'),

	)

def save_data(data1, n_clicks):
	if(n_clicks != None):
		df = pd.DataFrame(data1)
		#df.to_csv("sped_data.csv", index=False)
		return dcc.send_data_frame(df.to_csv, "checked_zeros.csv", index=False)

@callback(
	Output("download-dataframe-csv7f", "data"),
	
	Input('stored-data-8f', 'data'),
	Input('sped_button7f', 'n_clicks'),

	)

def save_data(data1, n_clicks):
	if(n_clicks != None):
		df = pd.DataFrame(data1)
		#df.to_csv("sped_data.csv", index=False)
		return dcc.send_data_frame(df.to_csv, "checked_zeros.csv", index=False)



@callback(
	Output("download-dataframe-csv2", "data"),
	Input('stored-data2', 'data'),
	Input('pal_button', 'n_clicks'),

	)

def save_data(data1, n_clicks):
	if(n_clicks != None):
		df = pd.DataFrame(data1)
		#df.to_csv("pal_data.csv", index=False)
		return dcc.send_data_frame(df.to_csv, "pal_data.csv", index=False)


@callback(
	Output("download-dataframe-csv2i", "data"),
	Input('stored-data2i', 'data'),
	Input('pal_buttoni', 'n_clicks'),

	)

def save_data(data1, n_clicks):
	if(n_clicks != None):
		df = pd.DataFrame(data1)
		#df.to_csv("pal_data.csv", index=False)
		return dcc.send_data_frame(df.to_csv, "pal_data_input.csv", index=False)

@callback(
	Output("download-dataframe-csv1", "data"),
	Input('stored-data3', 'data'),
	Input('ver_button', 'n_clicks'),

	)

def save_data(data1, n_clicks):
	if(n_clicks != None):
		df = pd.DataFrame(data1)
		return dcc.send_data_frame(df.to_csv, "ver_data.csv", index=False)


@callback(
	Output("download-dataframe-csv", "data"),
	Input('stored-data-3formated', 'data'),
	Input('btn', 'n_clicks'),
	)

def save_data(data1, n_clicks):
	if(n_clicks != None):
		df = pd.DataFrame(data1)
		return dcc.send_data_frame(df.to_csv, "filter_data.csv", index=False)


def german_to_float(val):
    if isinstance(val, str):
        val = val.strip()
        val = val.replace('.', '') 
        val = val[::-1].replace(',', '.', 1)[::-1]  # Replace last comma with dot
        return float(val)
    return val


########################################################
############# Log filter callback.  ####################
########################################################
@callback(
    Output('inflation-plot', 'figure'),
    Output('output-data-upload2', 'children'),
    Output('plot1', 'figure'),
    Output('plot2', 'figure'),
    Output('plot3', 'figure'),
    Output('plot4', 'figure'),
    Output('inflation-plot-4', 'figure'),
    Output('inflation-plot-5', 'figure'),
    Output('inflation-plot-6', 'figure'),
    Output('inflation-plot-10', 'figure'),
    Output('inflation-plot-11', 'figure'),

    Output('output-data-upload3', 'children'),
	Output('output-data-upload4', 'children'),
    Output('update-button', 'n_clicks'),
    Output('stored-data1', 'data'),
    Output('stored-data2', 'data'),
    Output('stored-data3', 'data'),
    Output('stored-data4', 'data'),
    Output('status', 'children'),
    

    Input('stored-data', 'data'),
    Input('update-button', 'n_clicks'),
    
)

def update_graph(data, n_clicks):

    default_return = ({}, html.Div(), {}, {}, {}, {}, {}, {}, {}, {}, {}, html.Div(), html.Div(), None, {}, {}, {}, [], "")

    df = pd.DataFrame(data)
    if n_clicks is None:
        if not df.empty:
        	default_return = ({}, html.Div(), {}, {}, {}, {}, {}, {}, {}, {}, {}, html.Div(), html.Div(), None, {}, {}, {}, [], "Daten erfolgreich geladen")
        return default_return


    if df.empty:
        default_return = ({}, html.Div(), {}, {}, {}, {}, {},{}, {}, {}, {}, html.Div(), html.Div(), None, {}, {}, {}, [], "Data Not loaded yet")
        return default_return

    if "Bestellmengeneinheit" in df.columns:
        default_return = ({}, html.Div(), {}, {}, {}, {}, {},{}, {}, {}, {}, html.Div(), html.Div(), None, {}, {}, {}, [], "Incorrect data")
        return default_return




    names_update_lower = {key.lower(): value for key, value in names_update.items()}
    df.columns = [names_update_lower.get(col.lower().strip(), col) for col in df.columns]

    print(df.columns)
    expected_columns = ["Auftragsmenge_Offen", "Auftragsmenge_bereits_geliefert", 
    "Summe von BrGew_Offen", "Fakturasperre","KomplettLF_KZ", "SalesOrder", "WE_PLZ", "Werk",
    "MatNr","BereitStellDat", "AME", "BME", "Zähler","MatBez","Auftragsmenge_Bestätigt","WE-Name","WE-Stadt","KzAZu"]
    missing_columns = [col for col in expected_columns if col not in df.columns]

    if(len(missing_columns) > 0):
        message = "Missing columns: " + ", ".join(missing_columns)
        default_return = ({}, html.Div(), {}, {}, {}, {}, {},{}, {}, {}, {}, html.Div(), html.Div(), None, {}, {}, {}, [], message)
        return default_return

    df['Auftragsmenge_Offen'] = df['Auftragsmenge_Offen'].apply(german_to_float)
    df['Auftragsmenge_bereits_geliefert'] = df['Auftragsmenge_bereits_geliefert'].apply(german_to_float)
    df['Summe von BrGew_Offen'] = df['Summe von BrGew_Offen'].apply(german_to_float)





    print(df["Summe von BrGew_Offen"].max())

    if "Summe von BrGew_Offen" not in df.columns:
        return default_return


    """cleaned_values = []
                to_keep = []
            
                for val in df["Summe von BrGew_Offen"]:
                    try:
                        val_str = str(val).strip()
                        val_str = val_str.replace(".", "", val_str.count(".") - 1).replace(",", ".")
                        cleaned_values.append(round(float(val_str),2))
                        to_keep.append(True)
                    except:
                        to_keep.append(False)
            
                df = df[to_keep].reset_index(drop=True)
                df["Summe von BrGew_Offen"] = cleaned_values"""

    if "Fakturasperre" in df.columns:
        df = df[df["Fakturasperre"].isna()].reset_index(drop=True)

    data_preprocessor2 = DataPreprocessing(df)
    data_req, ls, fig = data_preprocessor2.clac_2()
    ls2, ls3, data1, data2, fig2, fig3, fig4, fig5, fig6, fig7, fig10, fig20 = data_preprocessor2.get_calculated_results(input=False)
    fig8 = data_preprocessor2.get_absenders()

    return fig, ls, fig2, fig3, fig4, fig5, fig6, fig7, fig8, fig10, fig20,ls2, ls3, n_clicks, data_req, data1, data2, data, "Data Processed Successfully"




########################################################
############# Wareneingänge callback.  ####################
########################################################

@callback(
    Output('ploti1', 'figure'),
    Output('ploti2', 'figure'),
    Output('ploti3', 'figure'),
    Output('ploti4', 'figure'),
    Output('inflation-plot-10i', 'figure'),
    Output('inflation-plot-11i', 'figure'),
    Output('output-data-upload3i', 'children'),
    Output('stored-data2i', 'data'),
    Output('status20', 'children'),
    
    Input('stored-data-input', 'data'),
    Input('update-button2', 'n_clicks'),
    
)


def update_graph2(data, n_clicks):
	df = pd.DataFrame(data)
	if(n_clicks == None):
		if(df.empty):
			return {},{},{},{}, {}, {}, [], {}, "Daten erfolgreich not geladen"
		else:
			return {},{},{},{}, {}, {}, [], {}, "Daten erfolgreich geladen"
	else:

		if(df.empty):
			return {},{},{},{}, {}, {}, [], {}, "Daten erfolgreich not geladen"

		col_updates = {key.lower(): value for key, value in names_update.items()}
		df.columns = [col_updates.get(col.lower().strip(), col) for col in df.columns]


		#print(df.columns)
		expected_columns = ["Auftragsmenge_Offen", "Auftragsmenge_bereits_geliefert", 
		"Summe von BrGew_Offen", "Fakturasperre","KomplettLF_KZ", "SalesOrder", "WE_PLZ", "Werk",
		"MatNr","BereitStellDat", "AME", "BME", "Zähler","MatBez","Auftragsmenge_Bestätigt","WE-Name","WE-Stadt","KzAZu"]
		missing_columns = [col for col in expected_columns if col not in df.columns]

		if(len(missing_columns) > 0):
		    message = "Missing columns: " + ", ".join(missing_columns)
		    default_return = ({},{},{},{}, {}, {}, [], {},  message)
		    return default_return

		df['Auftragsmenge_Offen'] = df['Auftragsmenge_Offen'].apply(german_to_float)
		df['Auftragsmenge_bereits_geliefert'] = df['Auftragsmenge_bereits_geliefert'].apply(german_to_float)
		df['Summe von BrGew_Offen'] = df['Summe von BrGew_Offen'].apply(german_to_float)



		df = df[df["Fakturasperre"].isna()].reset_index(drop=True)

		data_preprocessor2 = DataPreprocessing(df)
		ls2, ls3, data1, data2, fig2, fig3, fig4, fig5, fig6, fig7, fig10, fig20 = data_preprocessor2.get_calculated_results(input=True)

		return fig2,fig3,fig4,fig5, fig10, fig20, ls2, data1,"Data Processed"


def format_to_int(value):
	formatted_value = value.replace('.', '').replace(',', '.')
	return int(float(formatted_value))



names_update = {
	"VStl": "Versandstelle",
	"Route": "Route",
	"Spediteur": "Forwarder",
	"VkOrg": "VKOrg",
	"VWeg": "VWEG",
	"SP": "Sparte",
	"VkBür": "VKBüro",
	"Angel.von": "SO_angelegt_von",
	"AuftrGeber": "AG-ID",
	"Warenempf.": "WE-ID",
	"Adresse": "WE-tatsächlich_ID",
	"Verursach.": "SalesOrder",
	"Eint": "SO_ATP_Einteilung",
	"AuftrMenge": "Auftragsmenge",
	"BestaMg": "Auftragsmenge_Bestätigt",
	"Offene Mng": "Auftragsmenge_Offen",
	#"AuftrMenge": "Auftragsmenge_Offen",
	"gelMenge": "Auftragsmenge_bereits_geliefert",
	"KMenge": "Auftragsmenge_bereits_geliefert",
	"ME": "AME",
	"BME": "BME",
	"KumAuMenge": "Auftragsmenge_Gesamt",
	"Werk": "Werk",
	"Kundenreferenz": "BestellNr_Kunde",
	"Versandbed": "Versandbedingung",
	"Zähler": "Zähler",
	"ZŠhler": "Zähler",
	"ZÃ¤hler": "Zähler",
	"SKU_Nenner": "SKU_Nenner",
	"Bereit.Dat": "BereitStellDat",
	"Name": "WE-Name",
	"Ort": "WE-Stadt",
	"LS": "Liefersperre",
	"Brutto": "Summe von BrGew_Offen",
	"KLF": "KomplettLF_KZ",
	"LfG": "LF-Gruppe",
	"VArt": "SO_Art",
	"Belegtyp": "BelegStatus",
	"Material": "MatNr",
	"Positionsbezeichnung": "MatBez",
	"Charge": "Charge_aus_SO",
	"Etyp": "ATP_Eint_Typ",
	"FS": "Fakturasperre",
	"IncTm": "Inco1",
	"Incoterms 2": "Inco2",
	"BDAr": "ATP_Bedarfsart",
	"Plz/WEMPF": "WE_PLZ"
}



@callback(
    Output('data_tab', 'children'),
    Output('stored-data-3formated', 'data'),
    Output('status2', 'children'),
 
    Input('stored-data-2', 'data'),
    Input('stored-data-3', 'data'),
    )
def update_second(input1, input2):

	df1 = pd.DataFrame(input1)
	df1 = df1.iloc[:, :4]

	data_dict = {}
	for i in range(len(df1)):
		id = df1["Unnamed: 1"][i]
		if(pd.isna(id)):
			continue
		else:
			if(id.isnumeric()):
				data = df1["Unnamed: 3"][i+1]
				data_dict[id] = data
	

	df2 = pd.DataFrame(input2)
	df2 = df2.iloc[3:]

	
	if(len(df2) > 1):
		data_check = dict(zip(df2["Unnamed: 0"], df2["Unnamed: 1"]))
		ls = []
		for key in data_check:
			if(key in data_dict.keys()):
				result = data_check[key] - format_to_int(data_dict[key])
				ls.append([key, data_check[key], format_to_int(data_dict[key]), result])

		data_df = pd.DataFrame(ls, columns=["Key", "input", "in system", "Differenz"])

		#print(data_df)

		ls = []
		ls.append(html.Div([
		html.H2("Tabelle",style={'color': 'black','font-size': '13px'}),
		dash_table.DataTable(
		    style_table={'height': '800px', 'overflowY': 'auto', 'width':'98%', 'margin-left':'4px'},
		    data=data_df.to_dict('records'),
		    columns=[{"name": i, "id": i} for i in data_df.columns],
		    
		    sort_action="native",
		    style_data={
            'backgroundColor': 'lightcyan',
            
        	},
		    style_header={
		        'backgroundColor': 'darkslategrey',
		        'color': 'white',
		        'fontWeight': 'bold',
		        'textAlign': 'center',
		        'border': '1px solid black'
		    }),html.Hr()])
		)

		#print(ls)

		return ls ,data_df.to_dict('records'), "data processed successfully"



	return html.Div([html.H2("Tabelle",style={'color': 'black','font-size': '13px'})]), {}, "data not uploaded"

import csv
def load_csv_with_best_encoding(file_bytes, encodings=None):
    if encodings is None:
        encodings = ['cp1252', 'ISO-8859-1', 'latin1', 'utf-8' ,'utf-8-sig', 'utf-16']
    delimiters = [';', ',']
    for enc in encodings:
        try:
            for sep in delimiters:
            	try:
            		df = pd.read_csv(io.BytesIO(file_bytes), encoding=enc, sep=sep)
            		if df.shape[1] > 1:
            			print(f"Successfully read with encoding='{enc}' and sep='{sep}'")
            			return df
            	except Exception:
            		continue

            
            # Optional: simple check for "garbled" characters in column names
            garbled = any('\ufffd' in col or '�' in col or any(ord(c) > 127 for c in col) for col in df.columns)
            if not garbled:
                print(f" Successfully loaded with encoding: {enc}")
                return df
            
            print(f"⚠ Loaded with encoding '{enc}', but column names may be garbled.")
            return df  # Return even if slightly garbled
        except Exception as e:
            print(f" Failed with encoding '{enc}': {e}")
    
    return pd.DataFrame()

def parse_contents(contents, filename, date):
    content_type, content_string = contents.split(',')

    print(filename)
    decoded = base64.b64decode(content_string)
    #print(filename)
    try:
        if 'csv' in filename:

            print(filename)
            # Assume that the user uploaded a CSV file
            df = load_csv_with_best_encoding(decoded, encodings=None)
            if df.empty:
            	df = pd.read_csv(io.StringIO(decoded.decode('utf-16')), delimiter='\t')

        if 'xls' in filename and not filename.endswith('xlsx'):
            df = pd.read_excel(io.BytesIO(decoded), engine='xlrd')
        elif filename.endswith('xlsx') or filename.endswith('XLSX'):
            df = pd.read_excel(io.BytesIO(decoded), engine='openpyxl')
    except Exception as e:
        print(e)
        return html.Div([
            'There was an error processing this file.'
        ])

    return html.Div([
        html.H5(children=filename,style={'text-align': 'center'}),
        #html.H6(datetime.datetime.fromtimestamp(date)),

        dash_table.DataTable(
            style_table={'height': '300px', 'overflowY': 'auto'},
            data=df.to_dict('records'),
            columns=[{"name": i, "id": i} for i in df.columns]
        ),

        html.Hr(),  # horizontal line
    ]), df

@callback(
    #Output('output-data-upload', 'children'),
    Output('stored-data', 'data'),
     #Output('status', 'children'),
    [Input('upload-data', 'contents')],
    [State('upload-data', 'filename'),
     State('upload-data', 'last_modified')]
)
def update_output(list_of_contents, list_of_names, list_of_dates):

    if list_of_contents is not None:
        children = []
        df_combined = pd.DataFrame()
        for c, n, d in zip(list_of_contents, list_of_names, list_of_dates):
            child, df = parse_contents(c, n, d)
            children.append(child)
            df_combined = pd.concat([df_combined, df], ignore_index=True)
        return df_combined.to_dict('records')#, #"Daten erfolgreich geladen"
    return {}#, ""

@callback(
    #Output('output-data-upload', 'children'),
    Output('stored-data-input', 'data'),
     
    [Input('upload-data2', 'contents')],
    [State('upload-data2', 'filename'),
     State('upload-data2', 'last_modified')]
)
def update_output20(list_of_contents, list_of_names, list_of_dates):

    if list_of_contents is not None:
        children = []
        df_combined = pd.DataFrame()
        for c, n, d in zip(list_of_contents, list_of_names, list_of_dates):
            child, df = parse_contents(c, n, d)
            children.append(child)
            df_combined = pd.concat([df_combined, df], ignore_index=True)
        return df_combined.to_dict('records')#, "Daten erfolgreich geladen"
    return {}#, ""

@callback(
    #Output('output-data-upload', 'children'),
    Output('stored-data-new', 'data'),
    Output('stored-data-old', 'data'),
    Output('statuson', 'children'),
    [Input('upload-newextracted', 'contents')],
    [Input('upload-existing', 'contents'),
    State('upload-newextracted', 'filename'),
     State('upload-newextracted', 'last_modified'), State('upload-existing', 'filename'),
     State('upload-existing', 'last_modified')]
)
def update_processedfiles(list_of_contents_new,list_of_contents_old, list_of_names_new, list_of_dates_new,list_of_names_old, list_of_dates_old):

    if list_of_contents_new is not None and list_of_contents_old is not None:
        children = []
        df_combined = pd.DataFrame()
        for c, n, d in zip(list_of_contents_new, list_of_names_new, list_of_dates_new):
            child, df = parse_contents(c, n, d)
            children.append(child)
            df_combined_new = pd.concat([df_combined, df], ignore_index=True)

        for c, n, d in zip(list_of_contents_old, list_of_names_old, list_of_dates_old):
            child, df = parse_contents(c, n, d)
            children.append(child)
            df_combined_old = pd.concat([df_combined, df], ignore_index=True)


        return df_combined_new.to_dict('records'), df_combined_old.to_dict('records'),"Daten erfolgreich geladen"

    elif(list_of_contents_old is not None and list_of_contents_new is None):
        return {},{},"Please upload new file as well"
    else:
        return {},{},""



@callback(
    #Output('output-data-upload', 'children'),
    Output('stored-data-5', 'data'),
    [Input('upload-data-3', 'contents')],
    [State('upload-data-3', 'filename'),
     State('upload-data-3', 'last_modified')]
)
def update_output3(list_of_contents, list_of_names, list_of_dates):

    if list_of_contents is not None:
        children = []
        df_combined = pd.DataFrame()
        for c, n, d in zip(list_of_contents, list_of_names, list_of_dates):
            child, df = parse_contents(c, n, d)
            children.append(child)
            df_combined = pd.concat([df_combined, df], ignore_index=True)
        return df_combined.to_dict('records')#, "Daten erfolgreich geladen"
    return {}#, ""

@callback(
    #Output('output-data-upload', 'children'),
    Output('stored-data-6', 'data'),
     Output('status6', 'children'),
    [Input('upload-data6', 'contents')],
    [State('upload-data6', 'filename'),
     State('upload-data6', 'last_modified')]
)
def update_output6(list_of_contents, list_of_names, list_of_dates):

    if list_of_contents is not None:
        children = []
        df_combined = pd.DataFrame()
        for c, n, d in zip(list_of_contents, list_of_names, list_of_dates):
            child, df = parse_contents(c, n, d)
            children.append(child)
            df_combined = pd.concat([df_combined, df], ignore_index=True)
        return df_combined.to_dict('records'), "Daten erfolgreich geladen"
    return {}, ""

@callback(
    #Output('output-data-upload', 'children'),
    Output('stored-data-7', 'data'),
     Output('status8', 'children'),
    [Input('upload-data7', 'contents')],
    [State('upload-data7', 'filename'),
     State('upload-data7', 'last_modified')]
)
def update_output8(list_of_contents, list_of_names, list_of_dates):

    if list_of_contents is not None:
        children = []
        df_combined = pd.DataFrame()
        for c, n, d in zip(list_of_contents, list_of_names, list_of_dates):
            child, df = parse_contents(c, n, d)
            children.append(child)
            df_combined = pd.concat([df_combined, df], ignore_index=True)
        return df_combined.to_dict('records'), "Daten erfolgreich geladen"
    return {}, ""



container_types = {
    1: "Drums",
    2: "Barrels",
    3: "Jerricans",
    4: "Boxes",
    5: "Bags",
    6: "Composite Packaging",
    7: "Pressure Receptical"
}

container_materials = {
    'A': "Steel",
    'B': "Aluminum",
    'C': "Natural Wood",
    'D': "Plywood",
    'F': "Reconstituted Wood",
    'G': "Fiberboard",
    'H': "Plastic",
    'L': "Textile",
    'M': "Paper",
    'N': "Metal other than Steel or Aluminum",
    'P': "Glass, Porcelain or Stoneware",
    'LQ': "Limited Quantity"
}


"""# Modified document processing function
def load_docx_and_print_tables(file_content, full_string, kgs, kgs2):
    try:
        # Create a file-like object from the uploaded content
        doc_file = io.BytesIO(file_content)
        document = Document(doc_file)

        # Process tables
        for table in document.tables:
            check = False
            count = 0
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.startswith("14"):
                        check = True
                    if check and cell.text.strip() == "":
                        count += 1
                        if count == 1 or count == 5:
                            continue
                        elif count == 2:
                            cell.text = full_string
                        elif count == 3:
                            cell.text = kgs
                        elif count == 4:
                            cell.text = kgs2

        # Save modified document to a BytesIO object
        modified_file = io.BytesIO()
        document.save(modified_file)
        modified_file.seek(0)
        return modified_file.getvalue()
    
    except Exception as e:
        print(f"Error processing document: {e}")
        return None"""

from docx.shared import Pt
def load_docx_and_print_tables(file_content,full_string, kgs, kgs2, target_row_number=None):
    try:
        # Load the document
        doc_file = io.BytesIO(file_content)
        document = Document(doc_file)


        # Check if there are any tables in the document
        if not document.tables:
            print("No tables found in this document.")
            return

        # Iterate through each table in the document
        
        #print(document.tables)
        #print("Tables:")
        
        table_count = 0
        for i, table in enumerate(document.tables):
            #print(f"\nTable {i + 1}:")  # Add 1 to i for user-friendly indexing
            check = False
            check_2 = False
            count = 0
            for row in table.rows:
                for cell in row.cells:
                    
                    if(cell.text.startswith("3")):
                        cell.text = ""#f"3. Page {table_count+1} of {table_count+1} pages"
                        paragraph = cell.paragraphs[0]
                        run = paragraph.add_run(f"3. Page {table_count+1} of {len(full_string)} pages")
                        run.font.size = Pt(8)  

                    if(cell.text.startswith("9") and check_2 == False):
                        paragraph = cell.paragraphs[3]
                        run = paragraph.add_run(f"\nBEFÖRDERUNG NACH ABSATZ 1.1.4.2.1")
                        run.font.size = Pt(10)
                        run.bold = True 
                        check_2 = True


                    #print(cell.text)
                    if(cell.text.startswith("14")):
                        
                        table_count +=1 
                        #print(table_count)
                        check = True
                    if(check == True and cell.text.strip() == ""):
                        #print(cell.text)
                        count+=1
                        if(count == 1 or count == 5):
                            continue
                        elif(count == 2):
                            cell.text = full_string[table_count-1]
                        elif(count == 3):
                            cell.text = kgs[table_count-1]
                        elif(count == 4):
                            cell.text = kgs2[table_count-1]
                            

            #print("-" * 20)  # Separator for readability
            
        modified_file = io.BytesIO()
        document.save(modified_file)
        modified_file.seek(0)
        return modified_file.getvalue()

    except Exception as e:
        print(f"Error processing document: {e}")
        return None

# Callback for Excel upload
@callback(
    Output('stored-data-72', 'data'),
    Output('status10', 'children'),
    Output('process-excel', 'disabled'),
    Input('upload-data72', 'contents'),
    State('upload-data72', 'filename'),
    State('upload-data72', 'last_modified')
)
def update_output8(contents, filename, date):
    if contents is not None:
        _, df = parse_contents(contents, filename, date)
        if not df.empty:
            return df.to_dict('records'), "Daten erfolgreich geladen: excel", False
    return {}, "Excel hochladen", True

# Callback for Word template upload
@callback(
    Output('stored-data-input_8', 'data'),
    Output('output-upload', 'children'),
    Input('upload-doc', 'contents'),
    State('upload-doc', 'filename')
)
def store_word_template(contents, filename):
    if contents is not None:
        return contents, html.Div([f"IMO geladen: {filename}"])
    return None, html.Div()

# Main processing callback
@callback(
    Output('status11', 'children'),
    Output("download-docx", "data"),
    Input('stored-data-72', 'data'),
    Input('stored-data-input_8', 'data'),
    Input('process-excel', 'n_clicks'),
    prevent_initial_call=True
)
def make_pdf_from_excel(data, word_template, n_clicks):
	if n_clicks is None or not data:
	    return "Bitte erst Daten hochladen ", None

	# Convert data to DataFrame
	excel_data = pd.DataFrame(data)

	if len(excel_data) == 0:
	    return "Keine validen Daten gefunden  ", None

	# Process Excel data
	full_string = ""
	kgs = ""
	kgs2 = ""
	count = 0
	list_of_strings = []
	list_of_kgs = []
	list_of_kgs2 = []


	for i in range(len(excel_data)):
	    
	    first_kg = str(excel_data["Brutto"][i])
	    first_kg_char = str(excel_data["Gewichtseinheit"][i])
	    
	    
	    second_kg = str(excel_data["Netto"][i])
	    second_kg_char = str(excel_data["Gewichtseinheit"][i])
	    
	    third = excel_data["Benennung"][i]
	    fourth = excel_data["UN-Nr."][i]
	    
	    if(pd.isna(fourth)):
	        continue
	    
	    get_data = excel_data["UN-Homologation"][i]
	    zero_word = excel_data["Menge"][i]
	    tech = excel_data["Tech.Benennung 1"][i]
	    
	    Flammpunkt = None
	    if("Flammpunkt" in excel_data.columns):
	        Flammpunkt = excel_data["Flammpunkt"][i]
	    
	    
	    if(pd.isna(get_data) or pd.isna(Flammpunkt)):
	        first_word = container_materials[get_data[1]]
	        second_word = container_types[int(get_data[0])]
	        data_to_add = str(zero_word)+" "+str(first_word)+" "+\
	                        str(second_word)+""+\
	                        str(fourth)+" "+str(third)+" "+str(tech) +"\n\n"
	        
	        full_string += data_to_add
	    else:
	        first_word = container_materials[get_data[1]]
	        second_word = container_types[int(get_data[0])]
	        data_to_add = str(zero_word)+" "+str(first_word)+" "+\
	                        str(second_word)+" ("+str(get_data)+")\n"+\
	                        str(fourth)+" "+str(third)+" "+str(tech) +"\n"+str(Flammpunkt)+"\n\n"
	        
	        
	        
	        full_string += data_to_add
	        
	            
	    
	    kgs += first_kg+" "+first_kg_char+"\n\n\n\n\n"
	    kgs2 += second_kg+" "+second_kg_char+"\n\n\n\n\n"
	    
	    
	    count += 1
	    if(count == 3 or i == len(excel_data)-1):
	        list_of_strings.append(full_string)
	        full_string = ""
	        list_of_kgs.append(kgs)
	        kgs = ""
	        list_of_kgs2.append(kgs2)
	        kgs2 = ""
	        count=0


	docx_file_path1 = 'layouts/IMO_Layout_new.docx'
	docx_file_path2 = 'layouts/IMO_Layout_2.docx'
	docx_file_path3 = 'layouts/IMO_Layout_3.docx'  
	docx_file_path4 = 'layouts/IMO_Layout_4.docx'  
	docx_file_path5 = 'layouts/IMO_Layout_5.docx'

	# Process Word document

	print("Length of list: ", len(list_of_strings))
	#print(list_of_strings)

	list_of_strings=[val for val in list_of_strings if val !=""]
	
	if word_template is None:
	    try:
	    	if(len(list_of_strings) == 1):
		        with open(docx_file_path1, "rb") as f:
		            word_content = f.read()
	    	if(len(list_of_strings) == 2):
		        with open(docx_file_path2, "rb") as f:
		            word_content = f.read()
	    	if(len(list_of_strings) == 3):
		        with open(docx_file_path3, "rb") as f:
		            word_content = f.read()
	    	if(len(list_of_strings) == 4):
		        with open(docx_file_path4, "rb") as f:
		            word_content = f.read()
	    	if(len(list_of_strings) == 5):
		        with open(docx_file_path5, "rb") as f:
		            word_content = f.read()
		            		            		            

	    except FileNotFoundError:
	        return "Kein Word Dokument verfügbar ", None
	else:
	    # Use uploaded template
	    content_type, content_string = word_template.split(',')
	    word_content = base64.b64decode(content_string)

	modified_docx_bytes = load_docx_and_print_tables(word_content, list_of_strings, list_of_kgs, list_of_kgs2)

	if modified_docx_bytes:
	    return "Dokument erfolgreich geladen", dcc.send_bytes(
	        modified_docx_bytes,
	        "IMO_Formular.docx"
	    )
	return "Error processing document", None





@callback(
    Output('status7', 'children'),
    Output('output-data-upload6', 'children'),
    Output('output-data-upload7', 'children'),
    Output('output-data-nonupdated', 'children'),
    Output('stored-data-6f', 'data'),
    Output('stored-data-7f', 'data'),
    Output('stored-data-8f', 'data'),
    

 	Input('stored-data-6', 'data'),
 	Input('stored-data-7', 'data'),
    )
def update_second(input1, input2):
	data = pd.DataFrame(input1)
	data2 = pd.DataFrame(input2)


	if(len(data) == 0):
		return "Data Not recived Yet", [], [], [], [],[],[]
	else:
		#print(data)
		data = data.dropna(subset=['Material','Charge','Gesperrt', 'Nicht freier Bestand', 'In Qualitätsprüfung' ])
		data = data.drop_duplicates(subset=['Charge'])


		selected_rows = data[(data['Gesperrt'] > 0) | (data['Nicht freier Bestand'] > 0) | (data['In Qualitätsprüfung'] > 0)]

		
		ls = []
		ls.append(html.Div([
		dash_table.DataTable(
		    style_table={'height': '400px','overflowY': 'auto', 'width':'98%', 'margin-left':'4px'},
		    data=selected_rows.to_dict('records'),
		    columns=[{"name": i, "id": i} for i in selected_rows.columns],
		    #editable=True,
		    #filter_action="native",
		    sort_action="native",
		    style_data={
            'backgroundColor': 'white',
            
        	},
		    #page_action="native",
		    style_header={
		        'backgroundColor': 'darkslategrey',
		        'color': 'lightcyan',
		        'fontWeight': 'bold',
		        'textAlign': 'center',
		        'border': '1px solid black'
		    }),html.Hr()])
		)

		new_dataframe = pd.DataFrame()
		new_dataframe2 = pd.DataFrame()

		if(len(data2) > 0):
			data2 = data2.dropna(subset=['Material','Charge','Gesperrt', 'Nicht freier Bestand', 'In Qualitätsprüfung' ]).reset_index(drop=True)
			
			checked = []
			for i in range(len(data2)):
			    mat_id = data2["Charge"][i]
			    
			    if(mat_id not in list(selected_rows.Charge)):
			        continue
			     
			    #print(mat_id)   
			    if(mat_id in checked):
			        continue
			    else:
			        checked.append(mat_id)   
			        if(data2["Gesperrt"][i] == 0 and data2["Nicht freier Bestand"][i]== 0 and data2["In Qualitätsprüfung"][i]== 0):
			            new_dataframe = pd.concat([new_dataframe, data2.iloc[[i]]], ignore_index=True)
			        else:
			        	new_dataframe2 = pd.concat([new_dataframe2, data2.iloc[[i]]], ignore_index=True)

		ls2 = []

		if(len(new_dataframe)>0):
			ls2.append(html.Div([
			dash_table.DataTable(
			    style_table={'height': '400px','overflowY': 'auto', 'width':'98%', 'margin-left':'4px'},
			    data=new_dataframe.to_dict('records'),
			    columns=[{"name": i, "id": i} for i in new_dataframe.columns],
			    #editable=True,
			    #filter_action="native",
			    sort_action="native",
			    style_data={
	            'backgroundColor': 'white',
	            
	        	},
			    #page_action="native",
			    style_header={
			        'backgroundColor': 'darkslategrey',
			        'color': 'lightcyan',
			        'fontWeight': 'bold',
			        'textAlign': 'center',
			        'border': '1px solid black'
			    }),html.Hr()])
			)
		ls3 = []
		if(len(new_dataframe2) > 0):
			ls3.append(html.Div([
			dash_table.DataTable(
			    style_table={'height': '400px','overflowY': 'auto', 'width':'98%', 'margin-left':'4px'},
			    data=new_dataframe2.to_dict('records'),
			    columns=[{"name": i, "id": i} for i in new_dataframe2.columns],
			    #editable=True,
			    #filter_action="native",
			    sort_action="native",
			    style_data={
	            'backgroundColor': 'white',
	            
	        	},
			    #page_action="native",
			    style_header={
			        'backgroundColor': 'darkslategrey',
			        'color': 'lightcyan',
			        'fontWeight': 'bold',
			        'textAlign': 'center',
			        'border': '1px solid black'
			    }),html.Hr()])
			)


		return "Data Recived", ls, ls2,ls3, selected_rows.to_dict('records'),new_dataframe.to_dict('records'),new_dataframe2.to_dict('records')



def process_data(data):
    column_values = []
    previous_value = ""
    for i, row in data.iterrows():
        if(i < 3):
            continue

        new_value = ""
        for j, val in enumerate(row):
            if(not pd.isna(val)):
                if(j == 5):
                    #print(j, val)
                    new_value = row[1]

        if(new_value == ""):
            column_values.append([previous_value, pd.to_datetime(row[8], format="%d.%m.%Y"),row[9], row[14]])

        else:
            previous_value = new_value 
    data_1 = pd.DataFrame(column_values,columns= ["order_id", "date", "value", "client"] ) 
    data_1 = data_1.dropna().reset_index(drop=True)
    ls = []
    for i, val in enumerate(data_1["value"]):
        ls.append(abs(float(str(val).strip().replace(",", "."))))


    data_1["value"] = ls
    
    
    
    return data_1

def make_barchart(order_id, proccessed_data):

    
    filtered_data = proccessed_data[proccessed_data['order_id'] == order_id]

    # Group by client and calculate the sum of value
    client_values = filtered_data.groupby('client')['value'].sum().reset_index()

    # Create an interactive bar chart using Plotly
    fig = px.bar(
        client_values,
        x='client',
        y='value',
        text='value',  # Display values on the bars
        labels={'client': 'Client ID', 'value': 'Total Value'},  # Customize axis labels
        title=f'Gesamte Bestellungen des Artikels {order_id} je Kunde ',
        #color_discrete_sequence=["#F5B323"]
    )

    # Customize appearance
    fig.update_traces(marker_color='#F5B323', textposition='outside')  # Bar color and text position
    fig.update_layout(
        xaxis=dict(title='KundenNr.', tickmode='linear'),
        yaxis=dict(title='Gesamtwert'),
        title=dict(font=dict(size=18), x=0.5),  # Center-align title
        template='plotly_white',  # Use a clean theme

    )

    # Show the figure
    return fig


import plotly.express as px

def make_linechart(order_id, proccessed_data):
    
    
    filtered_data = proccessed_data[proccessed_data['order_id'] == order_id]

    # Sort data by date to ensure proper line chart visualization
    filtered_data = filtered_data.sort_values(by='date')

    # Create a line chart using Plotly
    fig = px.line(
        filtered_data,
        x='date',
        y='value',
        labels={'date': 'Date', 'value': 'Value'},
        title=f'Zeitachse der Bestellungen des Artikels {order_id}',
        markers=True,  # Add markers to indicate data points
        color_discrete_sequence=["#F5B323"]
    )

    # Customize appearance
    fig.update_traces(line_color='#F5B323', marker=dict(size=8))
    fig.update_layout(
        xaxis=dict(title='Datum', tickformat='%d-%m-%y'),
        yaxis=dict(title='wert'),
        title=dict(font=dict(size=18), x=0.5),  # Center-align title
        template='plotly_white',
    )

    # Show the figure
    return fig

def make_dropdowns(lss):
	dropdown_options = []
	for value in lss:
		dropdown_options.append(
	    	{'label': value, 'value': value})
	return dropdown_options

def make_barchart_client(client_id, proccessed_data):

    
    filtered_data = proccessed_data[proccessed_data['client'] == client_id]

    # Group by client and calculate the sum of value
    client_values = filtered_data.groupby('order_id')['value'].sum().reset_index()

    # Create an interactive bar chart using Plotly
    fig = px.bar(
        client_values,
        x='order_id',
        y='value',
        text='value',  # Display values on the bars
        labels={'client': 'Client ID', 'value': 'Total Value'},  # Customize axis labels
        title=f' Gesamte Bestellungen des Kunden {client_id} je Artike',
        #color_discrete_sequence=["#F5B323"]
    )

    # Customize appearance
    fig.update_traces(marker_color='#F5B323', textposition='outside')  # Bar color and text position
    fig.update_layout(
        xaxis=dict(title='ArtikelNr.', tickmode='linear'),
        yaxis=dict(title='Gesamtwert'),
        title=dict(font=dict(size=18), x=0.5),  # Center-align title
        template='plotly_white',  # Use a clean theme

    )

    # Show the figure
    return fig


def make_linechart_client(client_id, proccessed_data):
    
    
    filtered_data = proccessed_data[proccessed_data['client'] == client_id]

    # Sort data by date to ensure proper line chart visualization
    filtered_data = filtered_data.sort_values(by='date')

    # Create a line chart using Plotly
    fig = px.line(
        filtered_data,
        x='date',
        y='value',
        labels={'date': 'Date', 'value': 'Value'},
        title=f'Gesamte Bestellungen des Kunden {client_id}',
        markers=True,  # Add markers to indicate data points,
       	#color_discrete_sequence=["#F5B323"]
    )

    # Customize appearance
    fig.update_traces(line_color='#F5B323', marker=dict(size=8))
    fig.update_layout(
        xaxis=dict(title='Datum', tickformat='%d-%m-%y'),
        yaxis=dict(title='Wert'),
        title=dict(font=dict(size=18), x=0.5),  # Center-align title
        template='plotly_white',
    )

    # Show the figure
    return fig

@callback(
    #Output('output-data-upload', 'children'),
   	Output('status4', 'children'),
   	Output('plot30', 'figure'),
   	Output('plot40', 'figure'),
   	Output('plot31', 'figure'),
   	Output('plot41', 'figure'),
   	Output('search-input9', 'options'),
   	Output('search-input10', 'options'),
   	#Output('status3', 'children'),
    Input('stored-data-5', 'data'),
    Input('search-input9', 'value'),
    Input('search-input10', 'value'),
    
) 
def get_new_Data(data,order_id, client_id):

	data = pd.DataFrame(data)

	if(data.empty):
		return (f"data not recived", {}, {},{},{}, [],[])

	if(len(data) > 0):

		print(data)


		row_1_values = data.iloc[1].astype(str).tolist()  

		expected_columns = ["Menge in ErfassME", "Kunde", "Buch.dat.","LOrt"]

		#expected_columns = ["order_id", "date", "value", "client"]
		missing_columns = [col for col in expected_columns if col not in row_1_values]

		if(len(missing_columns) > 0):
		    message = "Missing columns: " + ", ".join(missing_columns)
		    default_return = (f"data recived {message}", {}, {},{},{}, [],[])
		    return default_return


		proccessed_data = process_data(data)
		lss = set(proccessed_data.order_id)
		lss2 = set(proccessed_data.client)
		options =make_dropdowns(lss)
		options2 =make_dropdowns(lss2)
        
		if(order_id == "" or client_id == ""):
			return "wähle ArtikelNr. & KundenNr.", {}, {}, {}, {}, options, options2
		else:
			fig1 = make_barchart(order_id, proccessed_data)
			fig2 = make_linechart(order_id, proccessed_data)
			fig3 = make_barchart_client(client_id, proccessed_data)
			fig4 = make_linechart_client(client_id, proccessed_data)
			return "Erhaltene Daten zum verarbeiten", fig1, fig2, fig3, fig4, options,options2#,""
	else:
		return "data not recived yet for processing", {}, {},{},{}, [],[]#, ""





@callback(
    #Output('output-data-upload', 'children'),
    Output('stored-data-2', 'data'),
    Output('stored-data-3', 'data'),
    [Input('upload-data-2', 'contents')],
    [State('upload-data-2', 'filename'),
     State('upload-data-2', 'last_modified')]
)
def update_output2(list_of_contents, list_of_names, list_of_dates):

    if list_of_contents is not None:
        children = []
        df_combined = pd.DataFrame()
        for c, n, d in zip(list_of_contents, list_of_names, list_of_dates):
            child, df = parse_contents(c, n, d)
            
            children.append(df)
        
        return children[0].to_dict('records'),children[1].to_dict('records')
    return {}, {}


@callback(
    Output('plot7', 'figure'),  # Update this div with the output
    Input('submit-button', 'n_clicks'),  # Trigger on button click
    Input('text-input', 'value'),  # Get the value from the text input
    Input('stored-data4', 'data'),
)
def update_output(n_clicks, input_value,data):
	if n_clicks != None:  
		#print(input_value)
		df = pd.DataFrame(data)

		names_update_lower = {key.lower(): value for key, value in names_update.items()}
		#print(names_update_lower)
		df.columns = [names_update_lower.get(col.lower().strip(), col) for col in df.columns]

		#print(df)
		#print(input_value)

		df = df[df["Fakturasperre"].isna()].reset_index(drop=True)
		
		data_preprocessor2 = DataPreprocessing(df)
		fig =data_preprocessor2.make_customized_plot(input_value)
		return fig
	else:
		return {}


app = dash.Dash(__name__, external_stylesheets=external_stylesheets)
server = app.server


page_1_layout = html.Div(
	className="app-container", 

	children=[
	html.Div(style=nav_style, children=[nav_content]),
	dcc.Store(id='stored-data'),
	dcc.Store(id='stored-data1'),
	dcc.Store(id='stored-data2'),
	dcc.Store(id='stored-data2i'),
	dcc.Store(id='stored-data3'),
	dcc.Store(id='stored-data4'),
	


	#dascher etc
	html.Div(id='sped',children = [
		html.H2("Spedition Tabelle",style={'color': 'black','font-size': '13px'}),
		html.Button('Tabelle speichern', id='sped_button', n_clicks=None,style={'backgroundColor': '#F5B323','fontSize':"12px"}),
		],
		style={"height": "100%", 'width': '80%', 'float': 'right', 'backgroundColor': 'lightgray'}
		),

	html.Div(id='output-data-upload2',
		style={"height": "100%", 'width': '80%', 'float': 'right', 'backgroundColor': 'lightgray'}
		),

	html.Div(id='spedc',children = [
		html.H2("Spedition Chart",style={'color': 'black','font-size': '13px'}),
		],
		style={"height": "100%", 'width': '80%', 'float': 'right', 'backgroundColor': 'lightgray'}
		),
	dcc.Graph(id='inflation-plot',
		style={"height": "80%", 'width': '80%', 'float': 'right', 'backgroundColor': 'lightgray'}
		),


	#pal and picks tables
	html.Div(id='pal',children = [
		html.H2("Pal & PU Tabelle",style={'color': 'black','font-size': '13px'}),
		html.Button('Tabelle speichern', id='pal_button', n_clicks=None,style={'backgroundColor': '#F5B323','fontSize':"12px"}),
		],
		style={"height": "100%", 'width': '80%', 'float': 'right', 'backgroundColor': 'lightgray'}
		),
	html.Div(id='output-data-upload3',
		style={"height": "100%", 'width': '80%', 'float': 'right', 'backgroundColor': 'lightgray'}
		),
	html.Div(id='palc',children = [
		html.H2("Pal & PU Chart",style={'color': 'black','font-size': '13px'}),
		],
		style={"height": "100%", 'width': '80%', 'float': 'right', 'backgroundColor': 'lightgray'}
		),

	html.Div(
		id="basic_details2",
		className="outer-container1",
		children=[
		html.Div(
		className="plots-container",
		children=[
		
		# Plot 1
		html.Div(
		    className="left-container",
		    children=[
		        html.Div(
		            dcc.Graph(id='plot1',
		                #figure=fig3
		                ),
		            className="plot4",
		            style={"height": "100%", 'width': '40%', 'float': 'left','padding': '10px'}
		        ),]

		        ),

		html.Div(
		    className="left-container",
		    children=[
		        html.Div(
		            dcc.Graph(id='plot3',
		                #figure=fig4
		                ),
		            
		            style={"height": "100%", 'width': '60%', 'float': 'right','padding': '10px'}
		        ),]

		        ),
		])
		],
		style={"height": "50%", 'width': '80%', 'float': 'right', 'backgroundColor': 'lightgray'}
		),


	html.Div(
		id="basic_details3",
		className="outer-container1",
		children=[
		html.Div(
		className="plots-container",
		children=[
		
		# Plot 1
		html.Div(
		    className="left-container",
		    children=[
		        html.Div(
		            dcc.Graph(id='plot2',
		                #figure=fig3
		                ),
		            className="plot4",
		            style={"height": "100%", 'width': '40%', 'float': 'left','padding': '10px'}
		        ),]

		        ),

		html.Div(
		    className="left-container",
		    children=[
		        html.Div(
		            dcc.Graph(id='plot4',
		                #figure=fig4
		                ),
		            
		            style={"height": "100%", 'width': '60%', 'float': 'right','padding': '10px'}
		        ),]

		        ),
		])
		],
		style={"height": "50%", 'width': '80%', 'float': 'right', 'backgroundColor': 'lightgray'}
		),


	dcc.Download(id="download-dataframe-csv2"),
	dcc.Download(id="download-dataframe-csv3"),
	dcc.Download(id="download-dataframe-csv1"),

	##combined
	html.Div(id='ver',children = [
		html.H2("Vereint Data Tabelle",style={'color': 'black','font-size': '13px'}),
		html.Button('Tabelle speichern', id='ver_button', n_clicks=None,style={'backgroundColor': '#F5B323','fontSize':"12px"}),
		],
		style={"height": "100%", 'width': '80%', 'float': 'right', 'backgroundColor': 'lightgray'}
		),
	html.Div(id='output-data-upload4',
		style={"height": "100%", 'width': '80%', 'float': 'right', 'backgroundColor': 'lightgray'}
		),
	html.Div(id='verc',children = [
		html.H2("Vereint Data  Chart",style={'color': 'black','font-size': '13px'}),
		],
		style={"height": "100%", 'width': '80%', 'float': 'right', 'backgroundColor': 'lightgray'}
		),
	dcc.Graph(id='inflation-plot-4',
		style={"height": "80%", 'width': '80%', 'float': 'right', 'backgroundColor': 'lightgray'}
		),

	dcc.Graph(id='inflation-plot-5',
		style={"height": "80%", 'width': '80%', 'float': 'right', 'backgroundColor': 'lightgray'}
		),

	html.Div(id='neue',children = [
		html.H2("Neue Updated Charts",style={'color': 'black','font-size': '13px'}),
		dcc.Input(
		id='text-input', 
		type='text', 
		value='3022', 
		placeholder='Enter User Id'
		),
		html.Button(
		'Submit', 
		id='submit-button', 
		n_clicks=None,style={'backgroundColor': '#F5B323','fontSize':"12px"}
		),
		],
		style={"height": "100%", 'width': '80%', 'float': 'right', 'backgroundColor': 'lightgray'}
		),

	html.Div(
		id="basic_details4",
		className="outer-container1",
		children=[
		html.Div(
		className="plots-container",
		children=[
		
		# Plot 1
		html.Div(
		    className="left-container",
		    children=[
		        html.Div(
		            dcc.Graph(id='inflation-plot-6',
		                #figure=fig3
		                ),
		            className="plot4",
		            style={"height": "100%", 'width': '60%', 'float': 'left','padding': '10px'}
		        ),]

		        ),

		html.Div(
		    className="left-container",
		    children=[
		        html.Div(
		            dcc.Graph(id='plot7',
		                #figure=fig4
		                ),
		            
		            style={"height": "100%", 'width': '40%', 'float': 'right','padding': '10px'}
		        ),]

		        ),
		])
		],
		style={"height": "50%", 'width': '80%', 'float': 'right', 'backgroundColor': 'lightgray'}
		),

    html.Div(
        dcc.Graph(id='inflation-plot-10',
        style={"height": "50%", 'width': '80%', 'float': 'right', 'backgroundColor': 'lightgray'}
        ),

    ),

    html.Div(
        dcc.Graph(id='inflation-plot-11',
        style={"height": "50%", 'width': '80%', 'float': 'right', 'backgroundColor': 'lightgray'}
        ),

    ),

	


	    ],
	id="theme_change"

	)

page_2_layout = html.Div([
	html.Div(style=nav_style, children=[nav_content2]),
	dcc.Store(id='stored-data-2'),
	dcc.Store(id='stored-data-3'),
	dcc.Store(id='stored-data-3formated'),

	#html.H2(id= "idd",children =["Tabelle"],style={"height": "100%", 'width': '80%', 'float': 'right', 'backgroundColor': 'cadetblue'}),

	html.Div(id='data_tab',children = [
		],
		style={"height": "100%", 'width': '80%', 'float': 'right', 'backgroundColor': 'lightgray'}
		),

	html.Div(id= "This")
])

page_4_layout = html.Div([
	html.Div(style=nav_style, children=[nav_content3]),
	dcc.Store(id='stored-data-5'),

	html.Div(children = [
		html.H5("ArtikelNr. Charts"),
		],
		style={"height": "100%", 'width': '80%', 'float': 'right', 'backgroundColor': 'lightgray'}
		),
	html.Div(children = [
		html.Div(
		            dcc.Graph(id='plot30',
		                #figure=fig4
		                ),
		            style={"height": "100%", 'width': '100%', 'float': 'right','padding': '10px'}
		        ),
	
		],
		style={"height": "100%", 'width': '80%', 'float': 'right', 'backgroundColor': 'lightgray'}
		),
	html.Div(children = [
		html.Div(
		            dcc.Graph(id='plot40',
		                #figure=fig4
		                ),
		            
		            style={"height": "100%", 'width': '100%', 'float': 'right','padding': '10px'}
		        ),
	
		],
		style={"height": "100%", 'width': '80%', 'float': 'right', 'backgroundColor': 'lightgray'}
		),
	
	html.Div(children = [
		html.H5("KundenNr. Charts"),
		],
		style={"height": "100%", 'width': '80%', 'float': 'right', 'backgroundColor': 'lightgray'}
		),
	html.Div(children = [
		html.Div(
		            dcc.Graph(id='plot31',
		                #figure=fig4
		                ),
		            
		            style={"height": "100%", 'width': '100%', 'float': 'right','padding': '10px'}
		        ),
	
		],
		style={"height": "100%", 'width': '80%', 'float': 'right', 'backgroundColor': 'lightgray'}
		),
	html.Div(children = [
		html.Div(
		            dcc.Graph(id='plot41',
		                #figure=fig4
		                ),
		            
		            style={"height": "100%", 'width': '100%', 'float': 'right','padding': '10px'}
		        ),
	
		],
		style={"height": "100%", 'width': '80%', 'float': 'right', 'backgroundColor': 'lightgray'}
		),

	html.Div()
])



page_3_layout = html.Div([
	#html.Div(style=nav_style, children=[nav_content2]),


	html.Div(
	    [
	    	html.H1("Login", style={ "color":"black"}),
	        html.H6("ID", style={'font-size': '13px', "color":"black"}),
	        dcc.Input(id='input-text1', 
	                  type='text', 
	                  value='', 
	                  style={"width": "100%"}),

	        html.H6("Password", style={'font-size': '13px', "color":"black"}),
	        dcc.Input(id='input-text2', 
	                  type='password', 
	                  value='', 
	                  style={"width": "100%"}),

	        html.Br(),
	        dcc.Link('Login', href='/page-2',style={'color': 'firebrick','fontWeight': 'bold'}),
	        html.Div(id='success'),
	    ],
	    style={
	        "display": "flex",
	        "flexDirection": "column",
	        "alignItems": "center",
	        "justifyContent": "center",
	        "width": "300px",  # Set fixed width
	        "height": "300px",  # Set fixed height
	        "margin": "auto",  # Center horizontally
	        "position": "absolute",
	        "color":"white",
	        "top": "50%",
	        "left": "50%",
	        "transform": "translate(-50%, -50%)",
	        "border": "1px solid #ccc",
	        "padding": "20px",
	        "boxShadow": "0px 0px 10px rgba(0, 0, 0, 0.1)",
	        "borderRadius": "8px",
	        "backgroundColor": "#FFCC33"
	    }
	),

	#dcc.Link('New Filter', href='/page1',style={'color': 'yellow','fontWeight': 'bold'}),
	#dcc.Link('Go back to Logfilter', href='/page-2',style={'color': 'Yellow', 'fontWeight':'bold'}),

	#html.Div(id= "This")
])

import fitz  
import matplotlib.pyplot as plt
import re
from tqdm import tqdm



def get_punct_free(text):
    new_text = ""
    for char in text:
        if char not in string.punctuation:
            new_text += char
            
    return new_text


import google.generativeai as genai
from PIL import Image
from io import BytesIO
import base64
import os



def process_200_images(image_bytes_list):
    if len(image_bytes_list) > 80:
        return {"error": "The image_bytes_list must contain less than 110 images."}

    prompt_parts = []
    for index, image_bytes in enumerate(image_bytes_list):  # Enumerate to get index
        #print(image_bytes[0])
        try:
            image = Image.open(BytesIO(image_bytes[0]))
            buffered = BytesIO()
            image.save(buffered, format="JPEG")
            image_base64 = base64.b64encode(buffered.getvalue()).decode("utf-8")
            
            prompt_parts.append(f"Image {image_bytes[1]}:")
            prompt_parts.append({
                "mime_type": "image/jpeg",
                "data": image_base64,
            })
            
            
        except Exception as e:
            return {"error": f"Error processing image {index + 1}: {e}"}

    prompt_parts.append("""
    for given images with index end with "00" and "01" extract
    1) get image label
    2) Best before date:
    3) get ml value in digit after label 
    4) get artikel number a solo numeric number without any symbol or special character of length between 4 and 6
    
    for given image with index end with "02" extract written text
    for given image with index end with "03" extract written text

    example output:
    **Image PO: 18714618: 00**
    1) Sika® Primer-3 N
    2) 04/26
    3) 1000ml
    4) 122239

    **Image PO: 18714618: 01**
    1) Sika® Primer-3 N
    2) 04/26
    3) 1000ml
    4) 122239

    **Image PO: 18714618: 02**
    MAT 123489
    3009963434
    1039

    **Image PO: 18714618: 03**
    Sika® Primer-3 N
    1000ml
    MAT 123489
    3009963434
    719
    
    Return results for all the given images. 

    """)
    
    return prompt_parts


def check_apha(text):
    if re.search(r'[A-Za-z]', text):
        return True
    else:
        return False

@app.callback(
    Output("status9", "children"),
    Output("process-btn", "disabled"),
    Output("pdf-content", "data"),
    Output("pdf-processed", "data"),
    Output("pdf_results", "children"),
    Output("ploti12", "figure"),
    Output("stored_results", "data"),
    
    Input("process-btn", "n_clicks"),
    Input("upload-pdf", "contents"),
    State("pdf-content", "data"),
    State("pdf-processed", "data"),
    State("api_input", "value"),
    prevent_initial_call=True
)
def handle_pdf(n_clicks,upload_content, content, processed, api_input):
	triggered_id = ctx.triggered_id  

	if triggered_id == "upload-pdf":
		if upload_content is None:
		    return "", True, None, False  # No file uploaded, disable button
		return "Datei geladen", False, upload_content, False, [], {}, pd.DataFrame().to_dict('records')

	if triggered_id == "process-btn":
		if content is None:
		    return "Keine Datei geladen", dash.no_update, None, False, [], {},pd.DataFrame().to_dict('records')

		if processed:
		    return "PDF bereits verarbeitet ", True, content, True , [], {},data.DataFrame().to_dict('records')

		if content is not None:

			#print(contents)
			#print(api_input)

			os.environ["GEMINI_API_KEY"] = api_input 
			genai.configure(api_key=os.environ["GEMINI_API_KEY"]) 
			model = genai.GenerativeModel("gemini-1.5-flash")

			
			content_type, content_string = content[0].split(',')
			decoded = io.BytesIO(base64.b64decode(content_string))

			# Extract text from the PDF
			doc = fitz.open(stream=decoded, filetype="pdf")

			ref = {}
			multi_images = []
			images_list = []
			for page_number in range(len(doc)):
			     
			    page = doc[page_number]
			    text = page.get_text()
			    
			    text_data = text.split("\n")
			    
			    charge =[val.strip().lower() for val in text_data if len(val) > 1]
			    
			    second_val = ""
			    first_val = ""
			    if("flaschennummer" in charge[-1]):
			        second_val = charge[-1].split(" ")
			        if(len(second_val) >=1):
			            second_val = second_val[1]
			        else:
			            second_val = charge[charge.index("flaschennummer")+1]
			            
			    if("halb-charge" in charge):
			        first_val = charge[charge.index("halb-charge")+1]
			    

			        
			    text = ""
			    artikle = ""
			    
			    for i, tex in enumerate(text_data):
			        if(tex.startswith("PO")):
			            text=tex
			            
			        if(tex.startswith("FERT-Artikel-Nummer")):
			            artikle = text_data[i+1]
			        
			            
			    images = page.get_images(full=True) 
			    text_data = text.split(":")
			    
			    text = text.split('  ')[0].strip()
			    #print(len(text))
			    if(len(text) == 0):
			        #print(len(text))
			        continue
			        
			    images = page.get_images(full=True) 
			    for img_index, img in enumerate(images):
			        xref = img[0]  
			        
			        if(img_index>3):
			            break
			            
			        base_image = doc.extract_image(xref)  
			        image_bytes = base_image["image"]  
			 
			        images_list.append((image_bytes, str(text+": "+str(page_number)+str(img_index))))
			        #p = text+":"+str(page_number)+str(img_index)  
			        
			    ref[text.split(":")[1].strip()] = [first_val, second_val]
			    
			        
			    if((page_number != 0 and page_number%10 == 0) or page_number == len(doc)-1):
			        multi_images.append(images_list)
			        images_list = []


			df_dict = pd.DataFrame.from_dict(ref, orient='index', columns=['Charge', 'flaschennummer'])
			df_dict.index.name = 'Bild1_nummer'
			df_dict.reset_index(inplace=True)
  
			prompt_list =[]
			for bulk in multi_images:
			    prompt_parts = process_200_images(bulk)
			    prompt_list.append(prompt_parts)

			responses= []
			for prompt_parts in prompt_list:
			    try:
			        response = model.generate_content(prompt_parts)
			        responses.append(response)
			        print(f"response: {i} completed") 
			    except Exception as e:
			        print(e)

			page_results = {}
			for response in responses:
			    splited_respones = response.text.split("\n")
			    artikel_number = ""
			    img1_id = ""

			    for lines in splited_respones:
			        if(len(lines)>30):
			            continue
			        else:
			            if(lines.startswith("**")):
			                img1_id = lines.split(" ")[3][:-2]
			                img1_id = img1_id[-1]
			                artikel_number = lines.split(" ")[2][:-1]
			                if(artikel_number in page_results.keys()):
			                    continue
			                else:
			                    page_results[artikel_number] = [artikel_number]
			                    continue
			                
			            if(artikel_number in page_results.keys()):
			                if(lines ==""):
			                    continue
			                if(img1_id == "2"):
			                    if(len(lines)) == 10:
			                        if(not check_apha(lines)):
			                            print(artikel_number,"2",lines)
			                            page_results[artikel_number].append(lines)
			                    else:
			                        continue
			                    #print("2",lines)
			                if(img1_id == "3"):
			                    if(len(lines)) == 10:
			                        if(not check_apha(lines)):
			                            #print(artikel_number,"3",lines)
			                            page_results[artikel_number].append(lines)
			                    else:
			                        continue
			                
			                if(img1_id == "0" or img1_id == "1"):
			                    splited_lines =lines.split(")")
			                    if(len(splited_lines) > 1):
			                        #print(splited_lines[1])
			                        page_results[artikel_number].append(splited_lines[1])
			                    else:
			                        if(splited_lines[0] != ""):
			                            page_results[artikel_number].append(splited_lines[0])

			for val in page_results:
			    val_ = page_results[val]
			    val_  = val_[:11]
			    page_results[val] = val_
			data = pd.DataFrame(list(page_results.values()), columns=["Bild1_nummer", "Bild1_label", "Bild1_Date","Bild1 ML", "Artikel NO Bild1",
    "Bild2_label", "Bild2_Date" ,"Bild ML", "Artikel NO Bild2","BottleInfo1","BottleInfo2"
                                           ])

			df_merged = data.merge(df_dict, on='Bild1_nummer', how='inner')
            
			def create_match_column(df):

			    df['Match'] = 'nicht übereinstimmend'  
			    df.loc[
			        (df['Bild1_label'] == df['Bild2_label']) &
			        (df['Bild1 ML'] == df['Bild ML']) &
			        (df['BottleInfo1'] == df['BottleInfo2']) &
			        (df['Bild1_Date'] == df['Bild2_Date']) &
			        (df['Charge'] == df['BottleInfo1']) &
			        (df['Artikel NO Bild1'] == df['Artikel NO Bild2']),
			        'Match'
			    ] = 'übereinstimmend'

			    return df

			data =create_match_column(df_merged)

			ls2 = []
			ls2.append(html.Div([
				dash_table.DataTable(
				    style_table={'height': '400px','overflowY': 'auto', 'width':'98%', 'margin-left':'4px'},
				    data=data.to_dict('records'),
				    columns=[{"name": i, "id": i} for i in data.columns],
				    #editable=True,
				    filter_action="native",
				    sort_action="native",
				    style_data={
		            'backgroundColor': 'white',
		            
		        	},
				    #page_action="native",
				    style_header={
				        'backgroundColor': 'darkslategrey',
				        'color': 'lightcyan',
				        'fontWeight': 'bold',
				        'textAlign': 'center',
				        'border': '1px solid black'
				    }),html.Hr()])
				)

			result_counts = data['Match'].value_counts().reset_index()
			#result_counts.columns = ['Ergebnis', 'Anzahl']  # Rename columns for clarity

			# Create the bar chart using Plotly Express
			fig = px.bar(result_counts, x='Match', y='count', 
			             title='Summe der Ergebnisse',
			             color='Match',  
			             color_discrete_map={'übereinstimmend': '#F5B323', 'nicht übereinstimmend': 'black'}) 

			fig.update_layout(xaxis_title="Ergebnis", yaxis_title="Anzahl") 

			return  "PDF verarbeitet", True, content, True , ls2, fig,  data.to_dict('records')
	else:
		if upload_content is None:
		    return "Keine Datei geladen", True, None, False ,[], {}, pd.DataFrame().to_dict('records') # No file uploaded, disable button
		return "Datei geladen", False, upload_content, False, [], {}, pd.DataFrame().to_dict('records')

		#return html.Div(""), True, [], True, True, {},pd.DataFrame().to_dict('records')

# Callback for Excel upload
@callback(
    Output('pdf_results_after_check', 'children'),
    Input('stored-data-new', 'data'),
    State('stored-data-old', 'data'),
)
def update_tables(data_new, data_old):

	df_new =pd.DataFrame(data_new)
	df_old =pd.DataFrame(data_old)

	if(len(df_new)>0 and len(df_old)>0):

		df_new['Charge'] = df_new['Charge'].astype(str)
		df_new['flaschennummer'] = df_new['flaschennummer'].astype(float)

		df_old['Charge'] = df_old['Charge'].astype(str)
		df_old['flaschennummer'] = df_old['flaschennummer'].astype(float)

		matched = df_new.merge(df_old[['Charge', 'flaschennummer']], on=['Charge', 'flaschennummer'], how='inner')

		ls2 = []
		ls2.append(html.Div([
			dash_table.DataTable(
			    style_table={'height': '400px','overflowY': 'auto', 'width':'98%', 'margin-left':'4px'},
			    data=matched.to_dict('records'),
			    columns=[{"name": i, "id": i} for i in matched.columns],
			    #editable=True,
			    filter_action="native",
			    sort_action="native",
			    style_data={
			    'backgroundColor': 'white',
			    },
			    #page_action="native",
			    style_header={
			        'backgroundColor': 'darkslategrey',
			        'color': 'lightcyan',
			        'fontWeight': 'bold',
			        'textAlign': 'center',
			        'border': '1px solid black'
			    }),html.Hr()])
			)

		return ls2

	else:
		return []


# App layout
app.layout = html.Div([
	dcc.Location(id='url', refresh=False),
	#html.Div(id='page-content'),
	dcc.Download(id="download-dataframe-csv"),


	html.Div(id='page-1-content', style={'display': 'block'}, children=[
	    page_1_layout
	]),

	html.Div(id='page-2-content', style={'display': 'none'}, children=[
	    page_2_layout
	]),
	html.Div(id='page-4-content', style={'display': 'none'}, children=[
	    page_4_layout
	]),

	html.Div(id='page-3-content', style={'display': 'none'}, children=[
	    page_3_layout
	]),
	html.Div(id='page-5-content', style={'display': 'none'}, children=[
	    page_5_layout
	]),
	html.Div(id='page-6-content', style={'display': 'none'}, children=[
		dcc.Store(id='stored-data-6'),
		dcc.Store(id='stored-data-7'),
	    page_6_layout
	]),
	html.Div(id='page-7-content', style={'display': 'none'}, children=[
		dcc.Store(id='stored-data-8'),
		dcc.Store(id='stored-data-9'),
		dcc.Store(id='stored-data-new'),
		dcc.Store(id='stored-data-old'),
	    page_7_layout
	]),
	html.Div(id='page-10-content', style={'display': 'none'}, children=[
	    page_10_layout
	])

	

	])


@app.callback(
    [Output('page-1-content', 'style'),
     Output('page-2-content', 'style'),
     Output('page-3-content', 'style'),
     Output('page-4-content', 'style'),
     Output('page-5-content', 'style'),
     Output('page-6-content', 'style'),
     Output('page-7-content', 'style'),
     Output('page-10-content', 'style'),

     Output('success', 'children'),

     
     ],
    [Input('url', 'pathname'),
    Input('input-text1', 'value'),
    Input('input-text2', 'value'),
    ]
)
def display_page(pathname,id_, pass_):

    if(id_ == "log" and pass_ == "log"):#C3asar!

        if pathname == '/page-2':
            return {'display': 'block'}, {'display': 'none'} ,{'display': 'none'}, {'display': 'none'},{'display': 'none'},{'display': 'none'},{'display': 'none'},{'display': 'none'},""
        elif(pathname == "/page1"):
            return {'display': 'none'}, {'display': 'block'} ,{'display': 'none'},{'display': 'none'},{'display': 'none'}, {'display': 'none'},{'display': 'none'},{'display': 'none'},""
        elif(pathname == "/page-3"):
            return {'display': 'none'}, {'display': 'none'} ,{'display': 'none'}, {'display': 'block'},{'display': 'none'},{'display': 'none'},{'display': 'none'},{'display': 'none'},""
        elif(pathname == "/page_input"):
            return {'display': 'none'}, {'display': 'none'} ,{'display': 'none'}, {'display': 'none'},{'display': 'block'},{'display': 'none'},{'display': 'none'},{'display': 'none'},"" 
        elif(pathname == "/page_filter"):
            return {'display': 'none'}, {'display': 'none'} ,{'display': 'none'}, {'display': 'none'},{'display': 'none'},{'display': 'block'},{'display': 'none'},{'display': 'none'},"" 
        elif(pathname == "/image"):
            return {'display': 'none'}, {'display': 'none'} ,{'display': 'none'}, {'display': 'none'},{'display': 'none'},{'display': 'none'},{'display': 'block'},{'display': 'none'},"" 
        elif(pathname == "/excel"):
            return {'display': 'none'}, {'display': 'none'} ,{'display': 'none'}, {'display': 'none'},{'display': 'none'},{'display': 'none'},{'display': 'none'},{'display': 'block'},"" 
        else:
            return {'display': 'none'}, {'display': 'none'}, {'display': 'none'},{'display': 'none'},{'display': 'none'}, {'display': 'none'},{'display': 'none'},{'display': 'none'},""
    elif(id_ == "" and pass_ == ""):
        return {'display': 'none'}, {'display': 'none'}, {'display': 'block'},{'display': 'none'},{'display': 'none'},{'display': 'none'},{'display': 'none'},{'display': 'none'},html.H6("Bitte ID und Passwort eintragen",style={"color":"black"})
    else:
        return {'display': 'none'}, {'display': 'none'}, {'display': 'block'},{'display': 'none'},{'display': 'none'},{'display': 'none'},{'display': 'none'},{'display': 'none'},""


app.title = "LogFilter"


app.css.append_css({
    'external_url': 'https://maxcdn.bootstrapcdn.com/bootstrap/4.0.0/css/bootstrap.min.css'
})


if __name__ == '__main__':
	#app.run(debug=True)
    app.run(debug=True, port="8090")




