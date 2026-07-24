import pandas as pd
import plotly.express as px
from dash import Input, Output, html, dash_table, State,dcc
import dash
import fitz
from algorithms.pdf import calculate_pdf_scores
from algorithms.ocrllm import get_results
from algorithms.pal_pick_logic import get_results
import base64
import logging
from dash import ctx



def register_page_content_callbacks(app, data_manager):
    """Register callbacks that populate page content with uploaded data."""


    @app.callback(
        Output("pal_pick_data", "data"),
        Input("full", "n_clicks"),
        State("data_full", "data"),
        prevent_initial_call=True,
    )
    def download_full_data(n_clicks, data):

        if not data:
            raise dash.exceptions.PreventUpdate

        df = pd.DataFrame(data)

        return dcc.send_data_frame(
            df.to_csv,
            "pal_pick_results.csv",
            index=False
        )
    @app.callback(
        Output("results-table", "data"),
        Output("current-page", "data"),
        Input("next-btn", "n_clicks"),
        Input("prev-btn", "n_clicks"),
        Input("full", "n_clicks"),
        State("current-page", "data"),
        State("uploaded-cache-key", "data"),
        prevent_initial_call=True,
    )
    def change_page(next_clicks, prev_clicks, full,  page, cache_key):

        df = data_manager.get_dataframe(cache_key)

        if(df.empty):
            raise dash.exceptions.PreventUpdate

        trigger = ctx.triggered_id

        full = False
        if trigger == "next-btn":
            if (page + 1) * 10 < len(df):
                page += 1

        elif trigger == "prev-btn":
            if page > 0:
                page -= 1
        elif trigger == "full":
            start = 0
            end = len(df)
            full =True

        if not full:
            start = page * 10
            end = start + 10

        # Only these 10 rows are sent to the browser
        return df.iloc[start:end].to_dict("records"), page
    
    # Analysis Page Content
    @app.callback(
        Output('analysis-content', 'children'),
        Input("results-table", "data"),
        Input('url', 'pathname')
    )
    def update_analysis_content(input_data, pathname):
        if pathname != '/analysis' or not input_data:
            raise dash.exceptions.PreventUpdate
        
        if not input_data:
            return html.Div([
                html.Div([
                    html.Div("⚠️", style={'fontSize': '48px', 'marginBottom': '20px'}),
                    html.H3("Keine Daten verfügbar", style={
                        'color': 'rgb(75, 75, 75)',
                        'marginBottom': '10px'
                    }),
                    html.P("Bitte laden Sie eine PDF-Datei von der Startseite hoch.", style={
                        'color': 'rgb(124, 124, 124)',
                        'fontSize': '14px'
                    }),
                ], style={
                    'textAlign': 'center',
                    'padding': '60px 20px',
                    'backgroundColor': 'rgb(255, 255, 255)',
                    'borderRadius': '12px',
                    'border': '2px dashed rgb(208, 208, 208)',
                    'maxWidth': '500px',
                    'margin': '40px auto'
                })
            ])#, dash.no_update

        #logging.info(f"Path: {pathname}")
        #logging.info(f"Cache key: {cache_key}")

        df = pd.DataFrame(input_data)
        #df = data_manager.get_dataframe(cache_key)
        if df.empty:
            return html.Div([
                html.Div([
                    html.Div("❌", style={'fontSize': '48px', 'marginBottom': '20px'}),
                    html.H3("Data Not Found", style={
                        'color': 'rgb(75, 75, 75)',
                        'marginBottom': '10px'
                    }),
                    html.P("The data has expired or could not be loaded.", style={
                        'color': 'rgb(124, 124, 124)',
                        'fontSize': '14px'
                    })
                ], style={
                    'textAlign': 'center',
                    'padding': '60px 20px',
                    'backgroundColor': 'rgb(255, 255, 255)',
                    'borderRadius': '12px',
                    'border': '2px solid rgb(216, 40, 47)',
                    'maxWidth': '500px',
                    'margin': '40px auto'
                })
            ])#,dash.no_update
        
        CARD_STYLE = {
            "backgroundColor": "white",
            "borderRadius": "12px",
            "padding": "20px",
            "marginBottom": "25px",
            "boxShadow": "0 2px 8px rgba(0,0,0,0.08)",
            "border": "1px solid #e5e7eb",
        }
        HEADER_STYLE = {
            "background": "linear-gradient(135deg, rgb(110, 149, 178) 0%, rgb(80, 117, 141) 100%)",
            "padding": "30px",
            "borderRadius": "12px",
            "marginBottom": "25px",
            "boxShadow": "0 2px 8px rgba(0,0,0,0.08)",
        }
        # Generate analysis
        if(len(df) > 0):
            data_req,fig,fig2, fig3,fig4,fig5,fig6,fig7,fig8, fig10,fig20, data1, data2 = get_results(df)
            df = pd.DataFrame(data_req)
            df_half = pd.DataFrame(data1)
            df_full = pd.DataFrame(data2)

        return html.Div(
            [

                # Header
                html.Div(
                    [
                        html.H2("Datenanalyse"),
                    ],
                    style=HEADER_STYLE,
                ),

                # Table Card
                html.Div(
                    [
                        html.Div(
                        [
                            html.H4("Data Preview"),
                            dash_table.DataTable(
                                id='page-datatable_TOF',
                                columns=[{'name': c, 'id': c} for c in df.columns],
                                data=df.to_dict('records'),
                                page_current=0,
                                page_size=5,
                                page_action='native',
                                sort_action='native',
                                filter_action='native',
                                # ...all your existing style_* arguments...
                            ),
                        ],
                        style=CARD_STYLE,
                    )
                    ],
                    style=CARD_STYLE,
                ),

                # Graph 1 Card
                html.Div(
                    [
                        html.H4(""),
                        dcc.Graph(id="fig1", figure=fig),
                    ],
                    style=CARD_STYLE,
                ),

                # Table Card
                html.Div(
                    [
                        html.Div(
                        [
                            html.H4("Data Preview"),
                            html.Div(
                                dash_table.DataTable(
                                    id='page-datatable_palpicks',
                                    columns=[{'name': c, 'id': c} for c in df_half.columns],
                                    data=df_half.to_dict('records'),
                                    page_current=0,
                                    page_size=5,
                                    page_action='native',
                                    sort_action='native',
                                    filter_action='native',
                                    # ...all your existing style_* arguments...
                                ),
                                style={
                                "width": "100%",
                                "overflowX": "auto",
                            },
                            )
                            
                        ],
                        style=CARD_STYLE,
                    )
                    ],
                    style=CARD_STYLE,
                ),

                # Two Graph Cards
                html.Div(
                    [
                        html.Div(
                            [
                                html.H4(""),
                                dcc.Graph(id="fig2", figure=fig2),
                            ],
                            style={**CARD_STYLE, "flex": 1},
                        ),

                        html.Div(
                            [
                                html.H4(""),
                                dcc.Graph(id="fig3", figure=fig3),
                            ],
                            style={**CARD_STYLE, "flex": 1},
                        ),
                    ],
                    style={
                        "display": "flex",
                        "gap": "20px",
                    },
                ),

                # Two Graph Cards
                html.Div(
                    [
                        html.Div(
                            [
                                html.H4(""),
                                dcc.Graph(id="fig4", figure=fig4),
                            ],
                            style={**CARD_STYLE, "flex": 1},
                        ),

                        html.Div(
                            [
                                html.H4(""),
                                dcc.Graph(id="fig5", figure=fig5),
                            ],
                            style={**CARD_STYLE, "flex": 1},
                        ),
                    ],
                    style={
                        "display": "flex",
                        "gap": "20px",
                    },
                ),

                 # Table Card
                html.Div(
                    [
                        html.Div(
                        [
                            html.H4("Data Preview"),
                            html.Div(
                                dash_table.DataTable(
                                    id='page-datatable_palpickstof',
                                    columns=[{'name': c, 'id': c} for c in df_full.columns],
                                    data=df_full.to_dict('records'),
                                    page_current=0,
                                    page_size=5,
                                    page_action='native',
                                    sort_action='native',
                                    filter_action='native',
                                    # ...all your existing style_* arguments...
                                ),
                                style={
                                "width": "100%",
                                "overflowX": "auto",
                            },
                            )
                            
                        ],
                        style=CARD_STYLE,
                    )
                    ],
                    style=CARD_STYLE,
                ),
                # Two Graph Cards
                html.Div(
                    [
                        html.Div(
                            [
                                html.H4(""),
                                dcc.Graph(id="fig6", figure=fig6),
                            ],
                            style={**CARD_STYLE, "flex": 1},
                        ),

                        html.Div(
                            [
                                html.H4(""),
                                dcc.Graph(id="fig7", figure=fig7),
                            ],
                            style={**CARD_STYLE, "flex": 1},
                        ),
                    ],
                    style={
                        "display": "flex",
                        "gap": "20px",
                    },
                ),

                # Two Graph Cards
                html.Div(
                    [
                        html.Div(
                            [
                                html.H4(""),
                                dcc.Graph(id="fig10", figure=fig10),
                            ],
                            style={**CARD_STYLE, "flex": 1},
                        ),

                        html.Div(
                            [
                                html.H4(""),
                                dcc.Graph(id="fig20", figure=fig20),
                            ],
                            style={**CARD_STYLE, "flex": 1},
                        ),
                    ],
                    style={
                        "display": "flex",
                        "gap": "20px",
                    },
                ),

                # Two Graph Cards
                html.Div(
                    [
                        html.Div(
                            [
                                html.H4(""),
                                dcc.Graph(id="fig8", figure=fig8),
                            ],
                            style={**CARD_STYLE, "flex": 1},
                        ),
                    ],
                    style={
                        "display": "flex",
                        "gap": "20px",
                    },
                ),

            ],
            style={
                "padding": "25px",
                "backgroundColor": "#f4f6f9",
                "minHeight": "100vh",
            },
        )#,df_full.to_dict("records")
            
    
    @app.callback(
    Output("download-dataframe-pdf", "data"),
    Input('download', 'n_clicks'),
    State('stored_results', 'data'),
    prevent_initial_call=True
    )
    def save_data(n_clicks,data1):
        if not n_clicks:
            raise dash.exceptions.PreventUpdate
        if(n_clicks != None):
            df = pd.DataFrame(data1)
            #df.to_csv("sped_data.csv", index=False)
            return dcc.send_data_frame(df.to_csv, "pdf_details.csv", index=False)
        
        
    # Visualization Page Content
    @app.callback(
        Output('visualization-content', 'children'),
        Output('stored_results', 'data'),
        Input('viz-button', 'n_clicks'),  
        State('uploaded-cache-key', 'data'),
        State('url', 'pathname'),
        State('api_input', 'value'),
        State('model_selector', 'value'),
        prevent_initial_call=True
    )
    def update_visualization_content(n_clicks, cache_key, pathname, api_input, selected_model):

        print(selected_model)
        print(api_input)
        
        if not n_clicks:
            raise dash.exceptions.PreventUpdate
        
        if pathname != '/visualization':
            raise dash.exceptions.PreventUpdate
        if not cache_key:
            return html.Div([
                html.Div([
                    html.Div("📊", style={'fontSize': '48px', 'marginBottom': '20px'}),
                    html.H3("Keine Daten verfügbar", style={
                        'color': 'rgb(75, 75, 75)',
                        'marginBottom': '10px'
                    }),
                    html.P("Bitte laden Sie eine PDF-Datei von der Startseite hoch.", style={
                        'color': 'rgb(124, 124, 124)',
                        'fontSize': '14px'
                    })
                ], style={
                    'textAlign': 'center',
                    'padding': '60px 20px',
                    'backgroundColor': 'rgb(255, 255, 255)',
                    'borderRadius': '12px',
                    'border': '2px dashed rgb(208, 208, 208)',
                    'maxWidth': '500px',
                    'margin': '40px auto'
                })
            ]),pd.DataFrame().to_dict('records')
        
        print("Cache key for image processing: ", cache_key)
        obj = data_manager.get_data(cache_key)
        pdf_doc = fitz.open(stream=obj["bytes"], filetype="pdf")


       
        # Handle PDF
        if(selected_model == "ocrllm"):
            df = get_results(pdf_doc, selected_model,api_key=api_input)
        else:
            df = calculate_pdf_scores(pdf_doc,selected_model, api_key=api_input)

        correct = len(df[df['Match'] == 'übereinstimmend'])
        incorrect = len(df[df['Match'] == 'nicht übereinstimmend'])#(df['Match'] == '')

        if isinstance(obj, dict) and 'page_count' in obj:

            html_data = html.Div([
                    # Header
                    html.Div([
                        html.H2("PDF Ergebnisanalyse", style={
                            'color': 'rgb(255, 255, 255)',
                            'margin': '0',
                            'fontSize': '28px',
                            'fontWeight': '600'
                        }),
                        html.P(f"Dokument enthält {obj.get('page_count', '?')} Seiten | Größe: {obj.get('size_bytes', 0):,} bytes", style={
                            'color': 'rgba(255, 255, 255, 0.9)',
                            'margin': '8px 0 0 0',
                            'fontSize': '14px'
                        },),
                        html.P(f"nicht übereinstimmend: {incorrect} übereinstimmend: {correct}", style={
                            'color': 'rgba(255, 255, 255, 0.9)',
                            'margin': '8px 0 0 0',
                            'fontSize': '14px'
                        })

                    ], style={
                        'background': 'linear-gradient(135deg, rgb(110, 149, 178) 0%, rgb(80, 117, 141) 100%)',
                        'padding': '40px 20px',
                        'marginBottom': '30px'
                    }),
                    
                    # Data Table
                    dash_table.DataTable(
                        id='page-datatable_image',
                        columns=[{'name': c, 'id': c} for c in df.columns],
                        data=df.to_dict('records'),
                        page_current=0,
                        page_size=20,
                        page_action='native',
                        sort_action='native',
                        filter_action='native',
                        style_table={
                            'overflowX': 'auto',
                            'width': '100%'
                        },
                        style_cell={
                            'textAlign': 'left',
                            'minWidth': '100px',
                            'maxWidth': '300px',
                            'padding': '14px 16px',
                            'fontFamily': '-apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif',
                            'fontSize': '14px',
                            'border': 'none',
                            'borderBottom': '1px solid rgb(240, 240, 240)',
                            'whiteSpace': 'normal',
                            'height': 'auto',
                            'overflow': 'hidden',
                            'textOverflow': 'ellipsis'
                        },
                        style_header={
                            'backgroundColor': 'rgb(110, 149, 178)',
                            'color': 'white',
                            'fontWeight': '600',
                            'textTransform': 'uppercase',
                            'fontSize': '12px',
                            'letterSpacing': '0.5px',
                            'padding': '16px',
                            'border': 'none'
                        },
                        style_data={
                            'backgroundColor': 'rgb(255, 255, 255)',
                            'color': 'rgb(75, 75, 75)'
                        },
                        style_data_conditional=[
                            {
                                'if': {'row_index': 'odd'},
                                'backgroundColor': 'rgb(248, 249, 250)'
                            },
                            {
                                'if': {'state': 'active'},
                                'backgroundColor': 'rgba(110, 149, 178, 0.1)',
                                'border': '1px solid rgb(110, 149, 178)'
                            }
                        ],
                        tooltip_data=[
                            {
                                column: {'value': str(value), 'type': 'markdown'}
                                for column, value in row.items()
                            } for row in df.to_dict('records')
                        ],
                        tooltip_duration=None
                    )
                ], style={
                    'width': '100%',
                    'minHeight': '100vh',
                    'backgroundColor': 'rgb(248, 249, 250)',
                    'paddingBottom': '40px'
                })
            return html_data, df.to_dict('records')
        
        return html.Div(""), pd.DataFrame().to_dict('records')
    

    def calculate_stock_mdi(df):
        """
        Calculate stock weight for Storage Locations 4, 5, and 105.

        Returns:
            dict with individual sums, total, and status.
        """

        STORAGE_COL = "Plant"
        WEIGHT_COL = "MDI (%)"   # Column containing values like "20.000 kg"

        print(df.columns)

        # Copy dataframe

        print(df)
        data = df.copy()

        print(data)
        # Keep only required storage locations
        data = data[data[STORAGE_COL].isin([4, 5, 105])]

        # Extract numeric weight
        data["weight_kg"] = (
            data[WEIGHT_COL]
            .astype(str)
            .str.strip()
            .str.replace("\xa0", "", regex=False)
            .str.replace(".", "", regex=False)
            .str.replace(",", ".", regex=False)
        )

        data["weight_kg"] = pd.to_numeric(data["weight_kg"], errors="coerce")

        # Sum by storage location
        sums = data.groupby(STORAGE_COL)["weight_kg"].sum()

        #print(df["Storage Location"].unique())

        storage4 = sums.get(4, 0)
        storage5 = sums.get(5, 0)
        storage105 = sums.get(105, 0)

        total = storage4 + storage5 + storage105

        capacity = 200000.0          # 200 tonnes in kg
        green_limit = capacity * 0.80
        yellow_limit = capacity * 0.95

        if total < green_limit:
            status = "Green"
        elif total <= yellow_limit:
            status = "Yellow"
        else:
            status = "Red"

        return {
            "Lagerung 4 (kg)": round(storage4, 2),
            "Lagerung 5 (kg)": round(storage5, 2),
            "Lagerung 105 (kg)": round(storage105, 2),
            "Gesamt": round(total, 2),
            "Kapa": capacity,
            "Nutzungsgrad (%)": round(total / capacity * 100, 2),
            "Status": status
        }
    

    def stock_result_to_html(result):
        color_map = {
            "Green": "#28a745",
            "Yellow": "#ffc107",
            "Red": "#dc3545"
        }

        return html.Div(
            [
                html.H4("Resultat:"),

                html.Table(
                    [
                        html.Tr([html.Th("Lagerung 4"), html.Td(f"{result['Lagerung 4 (kg)']:,.2f} kg")]),
                        html.Tr([html.Th("Lagerung 5"), html.Td(f"{result['Lagerung 5 (kg)']:,.2f} kg")]),
                        html.Tr([html.Th("Lagerung 105"), html.Td(f"{result['Lagerung 105 (kg)']:,.2f} kg")]),
                        html.Tr([html.Th("Gesamt"), html.Td(f"{result['Gesamt']:,.2f} kg")]),
                        html.Tr([html.Th("Kapa"), html.Td(f"{result['Kapa']:,.2f} kg")]),
                        html.Tr([html.Th("Nutzungsgrad"), html.Td(f"{result['Nutzungsgrad (%)']:.2f}%")]),
                    ],
                    style={"width": "50%"}
                ),

                html.Br(),

                html.Div(
                    "....",
                    style={
                        "backgroundColor": color_map[result["Status"]],
                        "color": "white",
                        "padding": "12px",
                        "fontWeight": "bold",
                        "fontSize": "20px",
                        "textAlign": "center",
                        "borderRadius": "8px",
                        "width": "200px"
                    },
                ),
            ]
        )

    # Analysis Page Content
    @app.callback(
        Output('storage-content', 'children'),
        Input('url', 'pathname'),
        State('uploaded-cache-key', 'data')
    )
    def update_storage_content(pathname,cache_key):

        print(pathname, cache_key)
        if pathname != '/storage' or not cache_key:
            raise dash.exceptions.PreventUpdate
        
        if not cache_key:
            return html.Div([
                html.Div([
                    html.Div("⚠️", style={'fontSize': '48px', 'marginBottom': '20px'}),
                    html.H3("Keine Daten verfügbar", style={
                        'color': 'rgb(75, 75, 75)',
                        'marginBottom': '10px'
                    }),
                    html.P("Bitte laden Sie eine PDF-Datei von der Startseite hoch.", style={
                        'color': 'rgb(124, 124, 124)',
                        'fontSize': '14px'
                    }),
                ], style={
                    'textAlign': 'center',
                    'padding': '60px 20px',
                    'backgroundColor': 'rgb(255, 255, 255)',
                    'borderRadius': '12px',
                    'border': '2px dashed rgb(208, 208, 208)',
                    'maxWidth': '500px',
                    'margin': '40px auto'
                })
            ])#, dash.no_update

        #logging.info(f"Path: {pathname}")
        #logging.info(f"Cache key: {cache_key}")

        obj = data_manager.get_data(cache_key)

        df = pd.DataFrame(obj)
        
        if(not df.empty):
            #print(df)
            result = calculate_stock_mdi(df)

        return stock_result_to_html(result)
    
    # Table Page Content
    @app.callback(
        Output('table-page-content', 'children'),
        Output("download-docx", "data"),
        Input('download_word', "n_clicks"),
        State('uploaded-cache-key', 'data'),
        State('url', 'pathname')
    )
    def update_table_page_content(n_clicks,cache_key, pathname):

        if not cache_key:
            return html.Div([
                html.Div([
                    html.Div("📋", style={'fontSize': '48px', 'marginBottom': '20px'}),
                    html.H3("No Data Available", style={
                        'color': 'rgb(75, 75, 75)',
                        'marginBottom': '10px'
                    }),
                    html.P("Please upload a file from the home page to view data.", style={
                        'color': 'rgb(124, 124, 124)',
                        'fontSize': '14px'
                    })
                ], style={
                    'textAlign': 'center',
                    'padding': '60px 20px',
                    'backgroundColor': 'rgb(255, 255, 255)',
                    'borderRadius': '12px',
                    'border': '2px dashed rgb(208, 208, 208)',
                    'maxWidth': '500px',
                    'margin': '40px auto'
                })
            ]), None

        if pathname != '/ML':
            raise dash.exceptions.PreventUpdate
        
        
        obj = data_manager.get_data(cache_key)
        logging.info(f"Printing cache from inside ML app call {cache_key}")
        logging.info(f"TYPE: {type(obj)}")
        logging.info(f"CONTENT: {obj}")

        if obj is None:
            return html.Div([
                html.Div([
                    html.Div("❌", style={'fontSize': '48px', 'marginBottom': '20px'}),
                    html.H3("Data Not Found", style={
                        'color': 'rgb(75, 75, 75)',
                        'marginBottom': '10px'
                    }),
                    html.P("The data has expired or could not be loaded.", style={
                        'color': 'rgb(124, 124, 124)',
                        'fontSize': '14px'
                    })
                ], style={
                    'textAlign': 'center',
                    'padding': '60px 20px',
                    'backgroundColor': 'rgb(255, 255, 255)',
                    'borderRadius': '12px',
                    'border': '2px solid rgb(216, 40, 47)',
                    'maxWidth': '500px',
                    'margin': '40px auto'
                })
            ]), None
    
        
        # Handle DataFrame
        if isinstance(obj, pd.DataFrame):
            df = obj
            page_size = 10

            list_of_kgs, list_of_kgs2, list_of_strings = get_results_word(df)
            docx_file_path1 = 'wlayouts/IMO_Layout_new.docx'
            docx_file_path2 = 'wlayouts/IMO_Layout_2.docx'
            docx_file_path3 = 'wlayouts/IMO_Layout_3.docx'  
            docx_file_path4 = 'wlayouts/IMO_Layout_4.docx'  
            docx_file_path5 = 'wlayouts/IMO_Layout_5.docx'


            # Process Word document
            print("Length of list: ", len(list_of_strings))
            list_of_strings=[val for val in list_of_strings if val !=""]
            try:
                if(len(list_of_strings) == 1):
                    with open(docx_file_path1, "rb") as f:
                        word_content = f.read()
                if(len(list_of_strings) == 2):
                    with open(docx_file_path2, "rb") as f:
                        word_content = f.read()
                if(len(list_of_strings) == 3):
                    with open(docx_file_path3, "rb") as f:
                        word_content = f.read()
                if(len(list_of_strings) == 4):
                    with open(docx_file_path4, "rb") as f:
                        word_content = f.read()
                if(len(list_of_strings) == 5):
                    with open(docx_file_path5, "rb") as f:
                        word_content = f.read()                                                

            except FileNotFoundError:
                print("Kein Word Dokument verfügbar ")

            modified_docx_bytes = load_docx_and_print_tables(word_content, list_of_strings, list_of_kgs, list_of_kgs2)
            if not n_clicks:
                raise dash.exceptions.PreventUpdate
            if modified_docx_bytes:
                return html.Div([
                    # Header
                    html.Div([
                        html.H2("Data Table View", style={
                            'color': 'rgb(255, 255, 255)',
                            'margin': '0',
                            'fontSize': '28px',
                            'fontWeight': '600'
                        }),
                        html.P(f"Showing {len(df)} rows × {len(df.columns)} columns", style={
                            'color': 'rgba(255, 255, 255, 0.9)',
                            'margin': '8px 0 0 0',
                            'fontSize': '14px'
                        })
                    ], style={
                        'background': 'linear-gradient(135deg, rgb(110, 149, 178) 0%, rgb(80, 117, 141) 100%)',
                        'padding': '40px 20px',
                        'marginBottom': '30px'
                    }),
                    
                    # Data Table
                    dash_table.DataTable(
                        id='page-datatable',
                        columns=[{'name': c, 'id': c} for c in df.columns],
                        data=df.to_dict('records'),
                        page_current=0,
                        page_size=page_size,
                        page_action='native',
                        sort_action='native',
                        filter_action='native',
                        style_table={
                            'overflowX': 'auto',
                            'width': '100%'
                        },
                        style_cell={
                            'textAlign': 'left',
                            'minWidth': '100px',
                            'maxWidth': '300px',
                            'padding': '14px 16px',
                            'fontFamily': '-apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif',
                            'fontSize': '14px',
                            'border': 'none',
                            'borderBottom': '1px solid rgb(240, 240, 240)',
                            'whiteSpace': 'normal',
                            'height': 'auto',
                            'overflow': 'hidden',
                            'textOverflow': 'ellipsis'
                        },
                        style_header={
                            'backgroundColor': 'rgb(110, 149, 178)',
                            'color': 'white',
                            'fontWeight': '600',
                            'textTransform': 'uppercase',
                            'fontSize': '12px',
                            'letterSpacing': '0.5px',
                            'padding': '16px',
                            'border': 'none'
                        },
                        style_data={
                            'backgroundColor': 'rgb(255, 255, 255)',
                            'color': 'rgb(75, 75, 75)'
                        },
                        style_data_conditional=[
                            {
                                'if': {'row_index': 'odd'},
                                'backgroundColor': 'rgb(248, 249, 250)'
                            },
                            {
                                'if': {'state': 'active'},
                                'backgroundColor': 'rgba(110, 149, 178, 0.1)',
                                'border': '1px solid rgb(110, 149, 178)'
                            }
                        ],
                        tooltip_data=[
                            {
                                column: {'value': str(value), 'type': 'markdown'}
                                for column, value in row.items()
                            } for row in df.to_dict('records')
                        ],
                        tooltip_duration=None
                    )
                ], style={
                    'width': '100%',
                    'minHeight': '100vh',
                    'backgroundColor': 'rgb(248, 249, 250)',
                    'paddingBottom': '40px'
                }),dcc.send_bytes(
                    modified_docx_bytes,
                    "IMO_Formular.docx"
                )
        
        return html.Div("Unsupported data format.")
    


def get_results_word(excel_data):
    full_string = ""
    kgs = ""
    kgs2 = ""
    count = 0
    list_of_strings = []
    list_of_kgs = []
    list_of_kgs2 = []


    for i in range(len(excel_data)):

        first_kg = str(excel_data["Brutto"][i])
        first_kg_char = str(excel_data["Gewichtseinheit"][i])


        second_kg = str(excel_data["Netto"][i])
        second_kg_char = str(excel_data["Gewichtseinheit"][i])

        third = excel_data["Benennung"][i]
        fourth = excel_data["UN-Nr."][i]

        if(pd.isna(fourth)):
            
            if(count == 3 or i == len(excel_data)-1):
                list_of_strings.append(full_string)
                full_string = ""
                list_of_kgs.append(kgs)
                kgs = ""
                list_of_kgs2.append(kgs2)
                kgs2 = ""
                count=0

            continue

        get_data = excel_data["UN-Homologation"][i]
        zero_word = excel_data["Menge"][i]
        tech = excel_data["Tech.Benennung 1"][i]

        Flammpunkt = None
        if("Flammpunkt" in excel_data.columns):
            Flammpunkt = excel_data["Flammpunkt"][i]


        if(pd.isna(get_data) or pd.isna(Flammpunkt)):
            first_word = container_materials[get_data[1]]
            second_word = container_types[int(get_data[0])]
            data_to_add = str(zero_word)+" "+str(first_word)+" "+\
                            str(second_word)+""+\
                            str(fourth)+" "+str(third)+" "+str(tech) +"\n\n"

            full_string += data_to_add
        else:
            first_word = container_materials[get_data[1]]
            second_word = container_types[int(get_data[0])]
            data_to_add = str(zero_word)+" "+str(first_word)+" "+\
                            str(second_word)+" ("+str(get_data)+")\n"+\
                            str(fourth)+" "+str(third)+" "+str(tech) +"\n"+str(Flammpunkt)+"\n\n"



            full_string += data_to_add



        kgs += first_kg+" "+first_kg_char+"\n\n\n\n\n"
        kgs2 += second_kg+" "+second_kg_char+"\n\n\n\n\n"



        count += 1
        if(count == 3 or i == len(excel_data)-1):
            list_of_strings.append(full_string)
            full_string = ""
            list_of_kgs.append(kgs)
            kgs = ""
            list_of_kgs2.append(kgs2)
            kgs2 = ""
            count=0
    return list_of_kgs, list_of_kgs2, list_of_strings

from docx import Document
import io
container_types = {
    1: "Drums",
    2: "Barrels",
    3: "Jerricans",
    4: "Boxes",
    5: "Bags",
    6: "Composite Packaging",
    7: "Pressure Receptical"
}

container_materials = {
    'A': "Steel",
    'B': "Aluminum",
    'C': "Natural Wood",
    'D': "Plywood",
    'F': "Reconstituted Wood",
    'G': "Fiberboard",
    'H': "Plastic",
    'L': "Textile",
    'M': "Paper",
    'N': "Metal other than Steel or Aluminum",
    'P': "Glass, Porcelain or Stoneware",
    'LQ': "Limited Quantity"
}

from docx.shared import Pt
def load_docx_and_print_tables(file_content,full_string, kgs, kgs2, target_row_number=None):
    try:
        # Load the document
        doc_file = io.BytesIO(file_content)
        document = Document(doc_file)


        # Check if there are any tables in the document
        if not document.tables:
            print("No tables found in this document.")
            return

        # Iterate through each table in the document
        
        #print(document.tables)
        #print("Tables:")
        
        table_count = 0
        
        for i, table in enumerate(document.tables):
            #print(f"\nTable {i + 1}:")  # Add 1 to i for user-friendly indexing
            count = 0
            check = False
            check_2 = False
            for row in table.rows:
                for cell in row.cells:
                    
                    if(cell.text.startswith("3") and check == False):
                        cell.text = ""#f"3. Page {table_count+1} of {table_count+1} pages"
                        paragraph = cell.paragraphs[0]
                        run = paragraph.add_run(f"3. Page {table_count+1} of {len(full_string)} pages")
                        run.font.size = Pt(8)  

                    if(cell.text.startswith("9") and check_2 == False):
                        paragraph = cell.paragraphs[3]
                        run = paragraph.add_run(f"\nBEFÖRDERUNG NACH ABSATZ 1.1.4.2.1")
                        run.font.size = Pt(10)
                        run.bold = True 
                        check_2 = True


                    #print(cell.text)
                    if(cell.text.startswith("14")):
                        
                        table_count +=1 
                        #print(table_count)
                        check = True
                    if(check == True and cell.text.strip() == ""):
                        #print(cell.text)
                        count+=1
                        if(count == 1 or count == 5):
                            continue
                        elif(count == 2):
                            print(full_string[table_count-1])
                            print(cell)
                            cell.text = full_string[table_count-1]
                        elif(count == 3):
                            cell.text = kgs[table_count-1]
                        elif(count == 4):
                            cell.text = kgs2[table_count-1]
                            

            #print("-" * 20)  # Separator for readability
            
        modified_file = io.BytesIO()
        document.save(modified_file)
        modified_file.seek(0)
        return modified_file.getvalue()

    except Exception as e:
        print(f"Error processing document: {e}")
        return None