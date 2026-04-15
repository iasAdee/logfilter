from docx import Document
import pandas as pd
from docx.shared import Pt
import io

container_types = {
    1: "Drums",
    2: "Barrels",
    3: "Jerricans",
    4: "Boxes",
    5: "Bags",
    6: "Composite Packaging",
    7: "Pressure Receptical"
}

container_materials = {
    'A': "Steel",
    'B': "Aluminum",
    'C': "Natural Wood",
    'D': "Plywood",
    'F': "Reconstituted Wood",
    'G': "Fiberboard",
    'H': "Plastic",
    'L': "Textile",
    'M': "Paper",
    'N': "Metal other than Steel or Aluminum",
    'P': "Glass, Porcelain or Stoneware",
    'LQ': "Limited Quantity"
}


def load_docx_and_print_tables(file_content,full_string, kgs, kgs2, target_row_number=None):
    try:
        # Load the document
        doc_file = io.BytesIO(file_content)
        document = Document(doc_file)


        # Check if there are any tables in the document
        if not document.tables:
            print("No tables found in this document.")
            return

        # Iterate through each table in the document
        
        #print(document.tables)
        #print("Tables:")
        
        table_count = 0
        for i, table in enumerate(document.tables):
            #print(f"\nTable {i + 1}:")  # Add 1 to i for user-friendly indexing
            check = False
            check_2 = False
            count = 0
            for row in table.rows:
                for cell in row.cells:
                    
                    if(cell.text.startswith("3")):
                        cell.text = ""#f"3. Page {table_count+1} of {table_count+1} pages"
                        paragraph = cell.paragraphs[0]
                        run = paragraph.add_run(f"3. Page {table_count+1} of {len(full_string)} pages")
                        run.font.size = Pt(8)  

                    if(cell.text.startswith("9") and check_2 == False):
                        paragraph = cell.paragraphs[3]
                        run = paragraph.add_run(f"\nBEFÖRDERUNG NACH ABSATZ 1.1.4.2.1")
                        run.font.size = Pt(10)
                        run.bold = True 
                        check_2 = True


                    #print(cell.text)
                    if(cell.text.startswith("14")):
                        
                        table_count +=1 
                        #print(table_count)
                        check = True
                    if(check == True and cell.text.strip() == ""):
                        #print(cell.text)
                        count+=1
                        if(count == 1 or count == 5):
                            continue
                        elif(count == 2):
                            cell.text = full_string[table_count-1]
                        elif(count == 3):
                            cell.text = kgs[table_count-1]
                        elif(count == 4):
                            cell.text = kgs2[table_count-1]
                            

            #print("-" * 20)  # Separator for readability
            
        modified_file = io.BytesIO()
        document.save(modified_file)
        modified_file.seek(0)
        return modified_file.getvalue()

    except Exception as e:
        print(f"Error processing document: {e}")
        return None


def get_results_word(excel_data):
    full_string = ""
    kgs = ""
    kgs2 = ""
    count = 0
    list_of_strings = []
    list_of_kgs = []
    list_of_kgs2 = []


    for i in range(len(excel_data)):

        first_kg = str(excel_data["Brutto"][i])
        first_kg_char = str(excel_data["Gewichtseinheit"][i])


        second_kg = str(excel_data["Netto"][i])
        second_kg_char = str(excel_data["Gewichtseinheit"][i])

        third = excel_data["Benennung"][i]
        fourth = excel_data["UN-Nr."][i]

        if(pd.isna(fourth)):
            
            if(count == 3 or i == len(excel_data)-1):
                list_of_strings.append(full_string)
                full_string = ""
                list_of_kgs.append(kgs)
                kgs = ""
                list_of_kgs2.append(kgs2)
                kgs2 = ""
                count=0

            continue

        get_data = excel_data["UN-Homologation"][i]
        zero_word = excel_data["Menge"][i]
        tech = excel_data["Tech.Benennung 1"][i]

        Flammpunkt = None
        if("Flammpunkt" in excel_data.columns):
            Flammpunkt = excel_data["Flammpunkt"][i]


        if(pd.isna(get_data) or pd.isna(Flammpunkt)):
            first_word = container_materials[get_data[1]]
            second_word = container_types[int(get_data[0])]
            data_to_add = str(zero_word)+" "+str(first_word)+" "+\
                            str(second_word)+""+\
                            str(fourth)+" "+str(third)+" "+str(tech) +"\n\n"

            full_string += data_to_add
        else:
            first_word = container_materials[get_data[1]]
            second_word = container_types[int(get_data[0])]
            data_to_add = str(zero_word)+" "+str(first_word)+" "+\
                            str(second_word)+" ("+str(get_data)+")\n"+\
                            str(fourth)+" "+str(third)+" "+str(tech) +"\n"+str(Flammpunkt)+"\n\n"



            full_string += data_to_add



        kgs += first_kg+" "+first_kg_char+"\n\n\n\n\n"
        kgs2 += second_kg+" "+second_kg_char+"\n\n\n\n\n"



        count += 1
        if(count == 3 or i == len(excel_data)-1):
            list_of_strings.append(full_string)
            full_string = ""
            list_of_kgs.append(kgs)
            kgs = ""
            list_of_kgs2.append(kgs2)
            kgs2 = ""
            count=0
    return list_of_kgs, list_of_kgs2, list_of_strings